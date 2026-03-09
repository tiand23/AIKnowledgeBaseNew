"""
文档处理服务 - 处理Kafka消息，执行文件解析、向量化和索引
"""
import asyncio
import hashlib
import json
import re
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from sqlalchemy import select
from tika import parser as tika_parser

from app.clients.openai_chat_client import openai_chat_client
from app.clients.db_client import db_client
from app.clients.elasticsearch_client import es_client
from app.clients.kafka_client import kafka_client
from app.clients.minio_client import minio_client
from app.clients.redis_client import redis_client
from app.core.config import settings
from app.models.file import (
    DocumentVector,
    ChunkSource,
    FileUpload,
    ImageBlock,
    TableRow,
    FILE_STATUS_DONE,
    FILE_STATUS_FAILED,
    FILE_STATUS_PROCESSING,
)
from app.models.user import User
from app.services.embedding_service import embedding_service
from app.services.experience_service import experience_service
from app.services.profile_service import profile_service
from app.services.profile_service import ProfileStrategy
from app.services.relation_search_service import relation_search_service
from app.services.search_service import search_service
from app.utils.logger import get_logger

logger = get_logger(__name__)

try:
    import pdfplumber  # type: ignore
except Exception:
    pdfplumber = None  # type: ignore

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None  # type: ignore

try:
    import pytesseract  # type: ignore
except Exception:
    pytesseract = None  # type: ignore


class DocumentProcessorService:

    def __init__(self):
        self.chunk_size = 900
        self.chunk_overlap = 120
        self.chat_client = openai_chat_client
        self._local_inflight_keys: set[str] = set()
        self._local_inflight_lock = asyncio.Lock()
        self._processing_lock_ttl = max(60, int(settings.KAFKA_PROCESSING_LOCK_TTL_SEC))
        self._done_ttl = max(300, int(settings.KAFKA_IDEMPOTENCY_DONE_TTL_SEC))

    async def _mark_file_status(self, file_md5: str, user_id: int, status: int) -> None:
        if not db_client.SessionLocal:
            return

        try:
            async with db_client.SessionLocal() as db:
                result = await db.execute(
                    select(FileUpload).where(
                        FileUpload.file_md5 == file_md5,
                        FileUpload.user_id == user_id
                    )
                )
                file_record = result.scalar_one_or_none()
                if not file_record:
                    return

                file_record.status = status
                await db.commit()
        except Exception as e:
            logger.warning(f"更新文件状态失败: file_md5={file_md5}, status={status}, err={e}")

    @staticmethod
    def _processing_lock_key(file_md5: str, user_id: int) -> str:
        return f"kafka:doc:lock:{user_id}:{file_md5}"

    @staticmethod
    def _done_marker_key(file_md5: str, user_id: int) -> str:
        return f"kafka:doc:done:{user_id}:{file_md5}"

    async def _is_file_done(self, file_md5: str, user_id: int) -> bool:
        if not db_client.SessionLocal:
            return False
        try:
            async with db_client.SessionLocal() as db:
                result = await db.execute(
                    select(FileUpload.status).where(
                        FileUpload.file_md5 == file_md5,
                        FileUpload.user_id == user_id,
                    )
                )
                row = result.first()
                if not row:
                    return False
                return int(row[0]) == int(FILE_STATUS_DONE)
        except Exception as e:
            logger.warning(f"查询文件状态失败: file_md5={file_md5}, user_id={user_id}, err={e}")
            return False

    async def _acquire_processing_lock(self, lock_key: str, token: str) -> bool:
        redis = getattr(redis_client, "redis", None)
        if redis:
            try:
                acquired = await redis.set(lock_key, token, ex=self._processing_lock_ttl, nx=True)
                return bool(acquired)
            except Exception as e:
                logger.warning(f"Redis 获取处理锁失败，降级本地锁: key={lock_key}, err={e}")

        async with self._local_inflight_lock:
            if lock_key in self._local_inflight_keys:
                return False
            self._local_inflight_keys.add(lock_key)
            return True

    async def _release_processing_lock(self, lock_key: str, token: str) -> None:
        redis = getattr(redis_client, "redis", None)
        if redis:
            try:
                script = """
                if redis.call('GET', KEYS[1]) == ARGV[1] then
                    return redis.call('DEL', KEYS[1])
                end
                return 0
                """
                await redis.eval(script, 1, lock_key, token)
            except Exception as e:
                logger.warning(f"Redis 释放处理锁失败: key={lock_key}, err={e}")

        async with self._local_inflight_lock:
            self._local_inflight_keys.discard(lock_key)

    async def _set_done_marker(self, file_md5: str, user_id: int) -> None:
        redis = getattr(redis_client, "redis", None)
        if not redis:
            return
        try:
            await redis.set(self._done_marker_key(file_md5, user_id), "1", ex=self._done_ttl)
        except Exception as e:
            logger.warning(f"写入幂等完成标记失败: file_md5={file_md5}, user_id={user_id}, err={e}")

    async def _has_done_marker(self, file_md5: str, user_id: int) -> bool:
        redis = getattr(redis_client, "redis", None)
        if not redis:
            return False
        try:
            return bool(await redis.exists(self._done_marker_key(file_md5, user_id)))
        except Exception as e:
            logger.warning(f"读取幂等完成标记失败: file_md5={file_md5}, user_id={user_id}, err={e}")
            return False

    async def _publish_to_dlq(
        self,
        message_data: Dict[str, Any],
        message,
        fail_reason: str,
    ) -> bool:
        try:
            dlq_payload = {
                "event": "document_parse_failed",
                "fail_reason": self._short_text(fail_reason, 300),
                "source_topic": getattr(message, "topic", ""),
                "source_partition": getattr(message, "partition", -1),
                "source_offset": getattr(message, "offset", -1),
                "source_timestamp": getattr(message, "timestamp", None),
                "original_message": message_data,
            }
            topic = settings.KAFKA_DOCUMENT_PARSE_DLQ_TOPIC
            sent = await kafka_client.send_message(
                topic=topic,
                value=dlq_payload,
                key=str(message_data.get("file_md5") or ""),
            )
            if sent:
                logger.warning(
                    "消息已转入 DLQ: topic=%s, source_topic=%s, partition=%s, offset=%s, file_md5=%s",
                    topic,
                    getattr(message, "topic", ""),
                    getattr(message, "partition", -1),
                    getattr(message, "offset", -1),
                    message_data.get("file_md5"),
                )
            else:
                logger.error(
                    "DLQ 发送失败: topic=%s, source_topic=%s, partition=%s, offset=%s, file_md5=%s",
                    topic,
                    getattr(message, "topic", ""),
                    getattr(message, "partition", -1),
                    getattr(message, "offset", -1),
                    message_data.get("file_md5"),
                )
            return bool(sent)
        except Exception as e:
            logger.error(f"发送 DLQ 失败: {e}", exc_info=True)
            return False

    @staticmethod
    def _clean_text(value: Any) -> str:
        if value is None:
            return ""
        text = str(value).replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"(?<=[\u3040-\u30ff\u3400-\u9fff])\s+(?=[\u3040-\u30ff\u3400-\u9fff])", "", text)
        text = re.sub(r"(?<=[\u3040-\u30ff\u3400-\u9fff])\s*\|\s*(?=[\u3040-\u30ff\u3400-\u9fff])", "", text)
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
        return text.strip()

    @staticmethod
    def _short_text(value: Any, limit: int = 160) -> str:
        text = DocumentProcessorService._clean_text(value)
        if len(text) <= limit:
            return text
        return f"{text[:limit]}...(len={len(text)})"

    @staticmethod
    def _clean_preview_text(value: Any, source_type: Optional[str] = None) -> str:
        """
        证据预览文本清洗：
        - 去掉 table_header/table_row 中的 col_x 占位列噪声
        - 合并多余分隔符，保留可读信息
        """
        text = DocumentProcessorService._clean_text(value)
        if not text:
            return ""

        st = str(source_type or "").strip().lower()
        if "table" not in st and "row" not in st and "header" not in st:
            return text

        text = re.sub(r"(?:(?:^|\||;)\s*col_\d+\s*:?\s*)", " | ", text)
        text = re.sub(r"\|\s*\|+", "|", text)
        text = re.sub(r";\s*;+", ";", text)
        text = re.sub(r"(?:\s*\|\s*){2,}", " | ", text)
        text = re.sub(r"(?:\s*;\s*){2,}", " ; ", text)
        text = re.sub(r"^\s*[\|;:\-]+\s*|\s*[\|;:\-]+\s*$", "", text)
        text = DocumentProcessorService._clean_text(text)

        if not re.search(r"[\u3040-\u30ff\u3400-\u9fffA-Za-z0-9]", text):
            return ""
        return text

    @staticmethod
    def _is_vlm_parser(parser: Any) -> bool:
        name = str(parser or "").strip()
        return name in {"vlm_sheet_snapshot", "vlm_diagram"}

    def _resolve_vlm_image_path(self, src: Dict[str, Any], chunk_meta: Dict[str, Any]) -> Optional[str]:
        """
        VLM 来源必须绑定图片；优先 source 级 image_path，缺失时回退到 chunk 级 image_path。
        """
        src_path = self._clean_text(src.get("image_path"))
        if src_path:
            return src_path
        meta_path = self._clean_text(chunk_meta.get("image_path"))
        if meta_path:
            return meta_path
        return None

    def _parse_docx_blocks(self, file_data: bytes, file_name: str) -> List[Dict[str, Any]]:
        blocks: List[Dict[str, Any]] = []
        doc = Document(BytesIO(file_data))
        current_section = ""
        block_idx = 0

        for para in doc.paragraphs:
            text = self._clean_text(para.text)
            if not text:
                continue

            style_name = ""
            if para.style and para.style.name:
                style_name = para.style.name.lower()

            if style_name.startswith("heading"):
                current_section = text[:120]
                block_type = "heading"
            else:
                block_type = "paragraph"

            blocks.append(
                {
                    "block_index": block_idx,
                    "type": block_type,
                    "text": text,
                    "page": None,
                    "section": current_section or None,
                    "sheet": None,
                    "source_parser": "docx",
                    "file_type": "docx",
                    "file_name": file_name,
                }
            )
            block_idx += 1

        for table_idx, table in enumerate(doc.tables):
            header_cells: Optional[List[str]] = None
            for row_idx, row in enumerate(table.rows):
                cells = [self._clean_text(cell.text) for cell in row.cells]
                if not any(cells):
                    continue
                if header_cells is None:
                    header_cells = cells
                    block_type = "table_header"
                    text = " | ".join(cells)
                else:
                    pairs = []
                    for i, value in enumerate(cells):
                        col_name = header_cells[i] if i < len(header_cells) else f"col_{i+1}"
                        if value:
                            pairs.append(f"{col_name}: {value}")
                    text = " ; ".join(pairs) if pairs else " | ".join(cells)
                    block_type = "table_row"

                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": block_type,
                        "text": text,
                        "page": None,
                        "section": current_section or None,
                        "sheet": f"table_{table_idx+1}",
                        "table_name": f"table_{table_idx+1}",
                        "row_no": row_idx + 1,
                        "row_json": {f"col_{i+1}": v for i, v in enumerate(cells)},
                        "source_parser": "docx",
                        "file_type": "docx",
                        "file_name": file_name,
                    }
                )
                block_idx += 1

        return blocks

    def _detect_xlsx_header_row(self, sheet: Any, max_scan_rows: int = 25) -> Optional[int]:
        """
        启发式识别 xlsx 表头行，避免把“标题/说明行”误判为 header。
        返回 1-based 行号，识别不到返回 None。
        """
        header_hints = {
            "id", "no", "name", "title", "type", "status", "date", "sheet", "page",
            "項目", "項番", "番号", "名称", "名前", "摘要", "概要", "種別", "担当", "更新日", "作成日",
        }
        best_row: Optional[int] = None
        best_score = float("-inf")

        for row_no, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_no > max_scan_rows:
                break
            values = [self._clean_text(v) for v in row]
            non_empty = [v for v in values if v]
            if len(non_empty) < 2:
                continue

            textual_cnt = 0
            numeric_cnt = 0
            unique_cnt = len(set(non_empty))
            hint_hit = 0
            for cell in non_empty:
                low = cell.lower()
                if any(h in low for h in header_hints):
                    hint_hit += 1
                if re.search(r"[\u3040-\u30ff\u3400-\u9fffA-Za-z]", cell):
                    textual_cnt += 1
                elif re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?", cell):
                    numeric_cnt += 1

            score = (
                len(non_empty) * 1.6
                + unique_cnt * 0.5
                + textual_cnt * 1.2
                + hint_hit * 2.0
                - numeric_cnt * 0.9
            )
            if score > best_score:
                best_score = score
                best_row = row_no

        return best_row

    def _parse_xlsx_blocks(self, file_data: bytes, file_name: str) -> List[Dict[str, Any]]:
        blocks: List[Dict[str, Any]] = []
        wb = load_workbook(BytesIO(file_data), data_only=True, read_only=False)
        drawing_anchors_by_sheet = self._extract_drawing_anchors_by_sheet(file_data)
        drawing_graph_by_sheet = self._extract_xlsx_drawing_graph_by_sheet(file_data)
        block_idx = 0
        logger.info(
            "XLSX parse start: file=%s, sheets=%s, drawing_anchor_sheets=%s, drawing_graph_sheets=%s",
            file_name,
            len(wb.worksheets),
            len(drawing_anchors_by_sheet),
            len(drawing_graph_by_sheet),
        )

        for sheet_idx, sheet in enumerate(wb.worksheets, start=1):
            sheet_name = sheet.title
            header: Optional[List[str]] = None
            detected_header_row = self._detect_xlsx_header_row(sheet)
            row_no = 0
            data_rows = 0
            schedule_count = 0
            node_count = 0
            edge_count = 0
            sheet_block_start = block_idx

            for row_no, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = [self._clean_text(v) for v in row]
                if not any(values):
                    continue

                if header is None:
                    if detected_header_row and row_no < detected_header_row:
                        continue
                    header = [v or f"col_{i+1}" for i, v in enumerate(values)]
                    text = " | ".join(header)
                    block_type = "table_header"
                    row_json = {f"col_{i+1}": v for i, v in enumerate(values)}
                else:
                    pairs = []
                    row_json = {}
                    for i, value in enumerate(values):
                        if not value:
                            continue
                        col_name = header[i] if i < len(header) else f"col_{i+1}"
                        pairs.append(f"{col_name}: {value}")
                        row_json[col_name] = value
                    text = " ; ".join(pairs) if pairs else " | ".join(values)
                    block_type = "table_row"
                    data_rows += 1

                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": block_type,
                        "text": text,
                        "page": None,
                        "section": None,
                        "sheet": sheet_name,
                        "row": row_no,
                        "table_name": sheet_name,
                        "row_no": row_no,
                        "row_json": row_json,
                        "source_parser": "xlsx",
                        "file_type": "xlsx",
                        "file_name": file_name,
                    }
                )
                block_idx += 1

            schedule_rows = self._extract_schedule_rows_from_sheet(sheet)
            if not schedule_rows:
                anchors = drawing_anchors_by_sheet.get(sheet_idx, [])
                if anchors:
                    logger.info(
                        "XLSX schedule fallback by anchors: file=%s, sheet=%s, anchors=%s",
                        file_name,
                        sheet_name,
                        len(anchors),
                    )
                    schedule_rows = self._extract_schedule_rows_from_drawing_anchors(sheet, anchors)
            for item in schedule_rows:
                period_text = f"{item.get('start_label', '?')} -> {item.get('end_label', '?')}"
                task_text = item.get("task_name") or f"row_{item.get('row_no')}"
                role_hint = item.get("task_detail") or ""
                confidence = float(item.get("confidence") or 0.0)
                text = f"[schedule] sheet={sheet_name} ; task={task_text} ; period={period_text}"
                if role_hint:
                    text += f" ; detail={role_hint[:240]}"
                text += f" ; confidence={confidence:.2f}"

                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "schedule_row",
                        "text": text,
                        "page": None,
                        "section": None,
                        "sheet": sheet_name,
                        "row": item.get("row_no"),
                        "table_name": sheet_name,
                        "row_no": item.get("row_no"),
                        "row_json": {
                            "task": task_text,
                            "period_start": item.get("start_label"),
                            "period_end": item.get("end_label"),
                            "period_start_col": item.get("start_col"),
                            "period_end_col": item.get("end_col"),
                            "task_detail": role_hint,
                            "confidence": confidence,
                        },
                        "source_parser": "xlsx_schedule",
                        "file_type": "xlsx",
                        "file_name": file_name,
                    }
                )
                block_idx += 1
                schedule_count += 1

            graph_info = drawing_graph_by_sheet.get(sheet_idx, {})
            for node in graph_info.get("nodes", []):
                node_name = self._clean_text(node.get("name"))
                if not node_name:
                    continue
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "diagram_node",
                        "text": f"[diagram_node] {node_name}",
                        "page": None,
                        "section": None,
                        "sheet": sheet_name,
                        "row_json": {
                            "name": node_name,
                            "left": node.get("left"),
                            "top": node.get("top"),
                            "right": node.get("right"),
                            "bottom": node.get("bottom"),
                            "center_x": node.get("center_x"),
                            "center_y": node.get("center_y"),
                        },
                        "source_parser": "xlsx_diagram",
                        "file_type": "xlsx",
                        "file_name": file_name,
                    }
                )
                block_idx += 1
                node_count += 1

            for edge in graph_info.get("edges", []):
                src = self._clean_text(edge.get("src"))
                dst = self._clean_text(edge.get("dst"))
                if not src or not dst or src == dst:
                    continue
                rel_type = self._clean_text(edge.get("relation_type")) or "连接"
                confidence = float(edge.get("confidence") or 0.0)
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "diagram_edge",
                        "text": f"{src} -> {dst} ; rel={rel_type} ; confidence={confidence:.2f}",
                        "page": None,
                        "section": None,
                        "sheet": sheet_name,
                        "row_json": {
                            "src": src,
                            "dst": dst,
                            "relation_type": rel_type,
                            "confidence": confidence,
                        },
                        "source_parser": "xlsx_diagram",
                        "file_type": "xlsx",
                        "file_name": file_name,
                    }
                )
                block_idx += 1
                edge_count += 1

            logger.info(
                "XLSX sheet parsed: file=%s, sheet=%s, header_row=%s, rows_seen=%s, data_rows=%s, schedule_rows=%s, diagram_nodes=%s, diagram_edges=%s, blocks_added=%s",
                file_name,
                sheet_name,
                detected_header_row,
                row_no,
                data_rows,
                schedule_count,
                node_count,
                edge_count,
                block_idx - sheet_block_start,
            )

        logger.info("XLSX parse end: file=%s, total_blocks=%s", file_name, len(blocks))
        return blocks

    @staticmethod
    def _distance(p1: Tuple[float, float], p2: Tuple[float, float]) -> float:
        dx = p1[0] - p2[0]
        dy = p1[1] - p2[1]
        return (dx * dx + dy * dy) ** 0.5

    @staticmethod
    def _extract_anchor_grid(anchor: Any, ns_draw: Dict[str, str]) -> Optional[Tuple[int, int, int, int]]:
        fr = anchor.find("xdr:from", ns_draw)
        to = anchor.find("xdr:to", ns_draw)
        if fr is None:
            return None
        try:
            from_col = int(fr.findtext("xdr:col", default="0", namespaces=ns_draw) or "0") + 1
            from_row = int(fr.findtext("xdr:row", default="0", namespaces=ns_draw) or "0") + 1
            if to is not None:
                to_col = int(to.findtext("xdr:col", default=str(from_col), namespaces=ns_draw) or str(from_col)) + 1
                to_row = int(to.findtext("xdr:row", default=str(from_row), namespaces=ns_draw) or str(from_row)) + 1
            else:
                to_col = from_col + 1
                to_row = from_row + 1
        except Exception:
            return None
        left = min(from_col, to_col)
        right = max(from_col, to_col)
        top = min(from_row, to_row)
        bottom = max(from_row, to_row)
        return left, top, right, bottom

    def _extract_xlsx_drawing_graph_by_sheet(self, file_data: bytes) -> Dict[int, Dict[str, List[Dict[str, Any]]]]:
        """
        从 xlsx drawing XML 解析图形节点与连接边。
        返回：
        {
          sheet_index: {
            "nodes": [{"name"}...],
            "edges": [{"src","dst","relation_type","confidence"}...]
          }
        }
        """
        result: Dict[int, Dict[str, List[Dict[str, Any]]]] = {}
        try:
            zf = zipfile.ZipFile(BytesIO(file_data))
        except Exception:
            return result

        ns_rel = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
        ns_draw = {"xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"}
        ns_a = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}

        for sheet_index in range(1, 256):
            sheet_path = f"xl/worksheets/sheet{sheet_index}.xml"
            rel_path = f"xl/worksheets/_rels/sheet{sheet_index}.xml.rels"
            if sheet_path not in zf.namelist():
                break
            if rel_path not in zf.namelist():
                continue

            try:
                rel_root = ET.fromstring(zf.read(rel_path))
            except Exception:
                continue

            drawing_targets: List[str] = []
            for rel in rel_root.findall("rel:Relationship", ns_rel):
                rel_type = rel.attrib.get("Type", "")
                if not rel_type.endswith("/drawing"):
                    continue
                target = rel.attrib.get("Target", "")
                if not target:
                    continue
                if target.startswith("../"):
                    target = f"xl/{target[3:]}"
                elif target.startswith("/"):
                    target = target.lstrip("/")
                else:
                    target = f"xl/worksheets/{target}"
                drawing_targets.append(target)

            nodes: List[Dict[str, Any]] = []
            connectors: List[Dict[str, Any]] = []

            for drawing_xml in drawing_targets:
                if drawing_xml not in zf.namelist():
                    continue
                try:
                    draw_root = ET.fromstring(zf.read(drawing_xml))
                except Exception:
                    continue

                anchor_nodes = draw_root.findall("xdr:twoCellAnchor", ns_draw) + draw_root.findall("xdr:oneCellAnchor", ns_draw)
                for anchor in anchor_nodes:
                    grid = self._extract_anchor_grid(anchor, ns_draw)
                    if not grid:
                        continue
                    left, top, right, bottom = grid
                    center_x = (left + right) / 2.0
                    center_y = (top + bottom) / 2.0

                    sp = anchor.find("xdr:sp", ns_draw)
                    cxn = anchor.find("xdr:cxnSp", ns_draw)

                    if sp is not None:
                        texts = [self._clean_text(t.text) for t in sp.findall(".//a:t", ns_a) if self._clean_text(t.text)]
                        node_name = self._clean_text(" ".join(texts))
                        prst = ""
                        prst_node = sp.find(".//a:prstGeom", ns_a)
                        if prst_node is not None:
                            prst = (prst_node.attrib.get("prst") or "").lower()

                        if node_name and len(node_name) >= 2:
                            nodes.append(
                                {
                                    "name": node_name[:255],
                                    "left": float(left),
                                    "top": float(top),
                                    "right": float(right),
                                    "bottom": float(bottom),
                                    "center_x": center_x,
                                    "center_y": center_y,
                                }
                            )
                            continue

                        if prst in {
                            "line", "straightconnector1", "bentconnector2", "bentconnector3",
                            "bentconnector4", "bentconnector5", "curvedconnector2",
                            "curvedconnector3", "curvedconnector4", "curvedconnector5",
                            "rightarrow", "leftarrow", "uparrow", "downarrow", "leftrightarrow",
                        }:
                            connectors.append({"start": (float(left), float(top)), "end": (float(right), float(bottom))})
                            continue

                    if cxn is not None:
                        connectors.append({"start": (float(left), float(top)), "end": (float(right), float(bottom))})

            if not nodes:
                continue

            uniq_nodes: Dict[str, Dict[str, Any]] = {}
            for n in nodes:
                key = re.sub(r"\s+", "", (n.get("name") or "").lower())
                if not key:
                    continue
                if key not in uniq_nodes:
                    uniq_nodes[key] = n
            nodes = list(uniq_nodes.values())

            edges: List[Dict[str, Any]] = []
            if connectors and len(nodes) >= 2:
                seen = set()
                for conn in connectors:
                    start = conn["start"]
                    end = conn["end"]
                    src = min(nodes, key=lambda n: self._distance(start, (n["center_x"], n["center_y"])))
                    dst = min(nodes, key=lambda n: self._distance(end, (n["center_x"], n["center_y"])))
                    if src["name"] == dst["name"]:
                        continue
                    d_src = self._distance(start, (src["center_x"], src["center_y"]))
                    d_dst = self._distance(end, (dst["center_x"], dst["center_y"]))
                    if d_src > 25 or d_dst > 25:
                        continue
                    edge_key = (src["name"], dst["name"])
                    if edge_key in seen:
                        continue
                    seen.add(edge_key)
                    confidence = max(0.4, min(0.95, 1.0 - (d_src + d_dst) / 80.0))
                    edges.append(
                        {
                            "src": src["name"],
                            "dst": dst["name"],
                            "relation_type": "连接",
                            "confidence": round(confidence, 2),
                        }
                    )

            result[sheet_index] = {
                "nodes": [
                    {
                        "name": n["name"],
                        "left": n.get("left"),
                        "top": n.get("top"),
                        "right": n.get("right"),
                        "bottom": n.get("bottom"),
                        "center_x": n.get("center_x"),
                        "center_y": n.get("center_y"),
                    }
                    for n in nodes
                ],
                "edges": edges,
            }
            logger.info(
                "XLSX drawing图关系抽取: sheet=%s, nodes=%s, connectors=%s, edges=%s",
                sheet_index,
                len(nodes),
                len(connectors),
                len(edges),
            )

        return result

    def _extract_drawing_anchors_by_sheet(self, file_data: bytes) -> Dict[int, List[Dict[str, int]]]:
        """
        从 xlsx drawing XML 提取 shape 锚点。
        返回：{sheet_index(1-based): [{row_no, start_col, end_col}, ...]}
        """
        result: Dict[int, List[Dict[str, int]]] = {}
        try:
            zf = zipfile.ZipFile(BytesIO(file_data))
        except Exception:
            return result

        ns_rel = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
        ns_draw = {"xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"}

        for sheet_index in range(1, 256):
            sheet_path = f"xl/worksheets/sheet{sheet_index}.xml"
            rel_path = f"xl/worksheets/_rels/sheet{sheet_index}.xml.rels"
            if sheet_path not in zf.namelist():
                break
            if rel_path not in zf.namelist():
                continue

            try:
                rel_root = ET.fromstring(zf.read(rel_path))
            except Exception:
                continue

            drawing_targets: List[str] = []
            for rel in rel_root.findall("rel:Relationship", ns_rel):
                rel_type = rel.attrib.get("Type", "")
                if not rel_type.endswith("/drawing"):
                    continue
                target = rel.attrib.get("Target", "")
                if not target:
                    continue
                if target.startswith("../"):
                    target = f"xl/{target[3:]}"
                elif target.startswith("/"):
                    target = target.lstrip("/")
                else:
                    target = f"xl/worksheets/{target}"
                drawing_targets.append(target)

            anchors: List[Dict[str, int]] = []
            for drawing_xml in drawing_targets:
                if drawing_xml not in zf.namelist():
                    continue
                try:
                    draw_root = ET.fromstring(zf.read(drawing_xml))
                except Exception:
                    continue

                for anchor in draw_root.findall("xdr:twoCellAnchor", ns_draw):
                    fr = anchor.find("xdr:from", ns_draw)
                    to = anchor.find("xdr:to", ns_draw)
                    if fr is None or to is None:
                        continue
                    try:
                        from_col = int(fr.findtext("xdr:col", default="0", namespaces=ns_draw) or "0") + 1
                        to_col = int(to.findtext("xdr:col", default="0", namespaces=ns_draw) or "0") + 1
                        from_row = int(fr.findtext("xdr:row", default="0", namespaces=ns_draw) or "0") + 1
                        to_row = int(to.findtext("xdr:row", default="0", namespaces=ns_draw) or "0") + 1
                    except Exception:
                        continue

                    start_col = min(from_col, to_col)
                    end_col = max(from_col, to_col)
                    row_no = min(from_row, to_row)
                    if end_col <= start_col:
                        continue
                    anchors.append(
                        {
                            "row_no": row_no,
                            "start_col": start_col,
                            "end_col": end_col,
                        }
                    )
            if anchors:
                result[sheet_index] = anchors
                logger.info("XLSX drawing锚点抽取: sheet_index=%s, anchors=%s", sheet_index, len(anchors))

        return result

    def _extract_schedule_rows_from_drawing_anchors(
        self,
        sheet: Any,
        anchors: List[Dict[str, int]],
    ) -> List[Dict[str, Any]]:
        """
        兜底：从 drawing 锚点恢复甘特行区间（适用于条形图非单元格填充场景）。
        """
        result: List[Dict[str, Any]] = []
        if not anchors:
            return result

        max_row = min(int(getattr(sheet, "max_row", 0) or 0), 260)
        max_col = min(int(getattr(sheet, "max_column", 0) or 0), 220)
        if max_row <= 5 or max_col <= 5:
            return result

        header_limit = min(15, max_row)
        year_row = None
        numeric_rows: List[int] = []
        for r in range(1, header_limit + 1):
            nums = 0
            years = 0
            for c in range(1, max_col + 1):
                val = self._clean_text(sheet.cell(row=r, column=c).value)
                if re.fullmatch(r"\d{4}", val):
                    years += 1
                if re.fullmatch(r"\d{1,2}", val):
                    nums += 1
            if years >= 1 and year_row is None:
                year_row = r
            if nums >= 6:
                numeric_rows.append(r)
        if not numeric_rows:
            return result

        week_row = max(numeric_rows)
        month_candidates = [r for r in numeric_rows if r < week_row]
        month_row = month_candidates[-1] if month_candidates else None
        data_start_row = week_row + 1
        data_end_row = min(max_row, data_start_row + 180)

        timeline_cols: List[int] = []
        for c in range(1, max_col + 1):
            m_val = self._clean_text(sheet.cell(row=month_row, column=c).value) if month_row else ""
            w_val = self._clean_text(sheet.cell(row=week_row, column=c).value)
            if re.fullmatch(r"\d{1,2}", m_val) or re.fullmatch(r"\d{1,2}", w_val):
                timeline_cols.append(c)
        if len(timeline_cols) < 8:
            return result

        def ffill(values: List[str]) -> List[str]:
            out: List[str] = []
            last = ""
            for v in values:
                if v:
                    last = v
                out.append(last)
            return out

        year_vals = ffill(
            [self._clean_text(sheet.cell(row=year_row, column=c).value) if year_row else "" for c in timeline_cols]
        )
        month_vals = ffill(
            [self._clean_text(sheet.cell(row=month_row, column=c).value) if month_row else "" for c in timeline_cols]
        )
        week_vals = [self._clean_text(sheet.cell(row=week_row, column=c).value) for c in timeline_cols]

        col_label: Dict[int, str] = {}
        for i, col in enumerate(timeline_cols):
            y = year_vals[i] if i < len(year_vals) else ""
            m = month_vals[i] if i < len(month_vals) else ""
            w = week_vals[i] if i < len(week_vals) else ""
            parts = [p for p in [y, m and f"{m}月", w and f"W{w}"] if p]
            col_label[col] = "-".join(parts) if parts else f"C{col}"

        left_text_cols = [c for c in range(1, min(timeline_cols[0], 10))]
        if not left_text_cols:
            return result

        row_best: Dict[int, Dict[str, int]] = {}
        for a in anchors:
            row_no = int(a.get("row_no") or 0)
            start_col = int(a.get("start_col") or 0)
            end_col = int(a.get("end_col") or 0)
            if row_no < data_start_row or row_no > data_end_row:
                continue
            if end_col < timeline_cols[0] or start_col > timeline_cols[-1]:
                continue
            start_col = max(start_col, timeline_cols[0])
            end_col = min(end_col, timeline_cols[-1])
            if end_col <= start_col:
                continue
            span = end_col - start_col
            prev = row_best.get(row_no)
            if not prev or span > (prev["end_col"] - prev["start_col"]):
                row_best[row_no] = {"start_col": start_col, "end_col": end_col}

        for row_no, interval in sorted(row_best.items()):
            task_cells = [self._clean_text(sheet.cell(row=row_no, column=c).value) for c in left_text_cols]
            task_texts = [t for t in task_cells if t]
            if not task_texts:
                continue
            task_name = task_texts[0]
            task_detail = " / ".join(task_texts[1:4]) if len(task_texts) > 1 else ""
            start_col = interval["start_col"]
            end_col = interval["end_col"]
            result.append(
                {
                    "row_no": row_no,
                    "task_name": task_name,
                    "task_detail": task_detail,
                    "start_col": start_col,
                    "end_col": end_col,
                    "start_label": col_label.get(start_col, f"C{start_col}"),
                    "end_label": col_label.get(end_col, f"C{end_col}"),
                    "confidence": 0.82,
                }
            )

        logger.info(
            "XLSX drawing日程抽取: sheet=%s, anchors=%s, rows=%s",
            getattr(sheet, "title", "-"),
            len(anchors),
            len(result),
        )
        return result

    @staticmethod
    def _is_visually_filled(cell: Any) -> bool:
        try:
            fill = getattr(cell, "fill", None)
            if not fill:
                return False
            if not getattr(fill, "fill_type", None):
                return False

            fg = getattr(fill, "fgColor", None)
            if not fg:
                return False
            rgb = (getattr(fg, "rgb", None) or "").upper()
            indexed = getattr(fg, "indexed", None)

            if rgb in ("", "00000000", "00FFFFFF", "FFFFFFFF"):
                return False
            if indexed in (0, 64):
                return False
            return True
        except Exception:
            return False

    @staticmethod
    def _cell_fill_signature(cell: Any) -> str:
        try:
            fill = getattr(cell, "fill", None)
            if not fill or not getattr(fill, "fill_type", None):
                return ""
            fg = getattr(fill, "fgColor", None)
            if not fg:
                return ""
            rgb = (getattr(fg, "rgb", None) or "").upper()
            indexed = getattr(fg, "indexed", None)
            theme = getattr(fg, "theme", None)
            tint = getattr(fg, "tint", None)
            return f"{fill.fill_type}|{rgb}|{indexed}|{theme}|{tint}"
        except Exception:
            return ""

    def _extract_schedule_rows_from_sheet(self, sheet: Any) -> List[Dict[str, Any]]:
        """
        从 Excel sheet 中抽取甘特类日程行。
        规则：
        1) 自动识别 timeline 列（头部有 year/month/week 数字）
        2) 通过填充色区间推断开始/结束
        """
        result: List[Dict[str, Any]] = []
        max_row = min(int(getattr(sheet, "max_row", 0) or 0), 260)
        max_col = min(int(getattr(sheet, "max_column", 0) or 0), 220)
        if max_row <= 5 or max_col <= 5:
            return result

        header_limit = min(15, max_row)
        year_row = None
        numeric_rows: List[int] = []
        for r in range(1, header_limit + 1):
            nums = 0
            years = 0
            for c in range(1, max_col + 1):
                val = self._clean_text(sheet.cell(row=r, column=c).value)
                if re.fullmatch(r"\d{4}", val):
                    years += 1
                if re.fullmatch(r"\d{1,2}", val):
                    nums += 1
            if years >= 1 and year_row is None:
                year_row = r
            if nums >= 6:
                numeric_rows.append(r)

        if not numeric_rows:
            return result
        week_row = max(numeric_rows)
        month_candidates = [r for r in numeric_rows if r < week_row]
        month_row = month_candidates[-1] if month_candidates else None
        data_start_row = week_row + 1

        timeline_cols: List[int] = []
        for c in range(1, max_col + 1):
            m_val = self._clean_text(sheet.cell(row=month_row, column=c).value) if month_row else ""
            w_val = self._clean_text(sheet.cell(row=week_row, column=c).value)
            if re.fullmatch(r"\d{1,2}", m_val) or re.fullmatch(r"\d{1,2}", w_val):
                timeline_cols.append(c)
        if len(timeline_cols) < 8:
            return result

        def ffill(values: List[str]) -> List[str]:
            out: List[str] = []
            last = ""
            for v in values:
                if v:
                    last = v
                out.append(last)
            return out

        year_vals = ffill(
            [self._clean_text(sheet.cell(row=year_row, column=c).value) if year_row else "" for c in timeline_cols]
        )
        month_vals = ffill(
            [self._clean_text(sheet.cell(row=month_row, column=c).value) if month_row else "" for c in timeline_cols]
        )
        week_vals = [self._clean_text(sheet.cell(row=week_row, column=c).value) for c in timeline_cols]

        col_label: Dict[int, str] = {}
        for i, col in enumerate(timeline_cols):
            y = year_vals[i] if i < len(year_vals) else ""
            m = month_vals[i] if i < len(month_vals) else ""
            w = week_vals[i] if i < len(week_vals) else ""
            parts = [p for p in [y, m and f"{m}月", w and f"W{w}"] if p]
            col_label[col] = "-".join(parts) if parts else f"C{col}"

        left_text_cols = [c for c in range(1, min(timeline_cols[0], 10))]
        if not left_text_cols:
            return result

        fill_counter: Counter = Counter()
        data_end_row = min(max_row, data_start_row + 180)
        timeline_total_cells = 0
        for r in range(data_start_row, data_end_row + 1):
            task_cells = [self._clean_text(sheet.cell(row=r, column=c).value) for c in left_text_cols]
            if not any(task_cells):
                continue
            for c in timeline_cols:
                timeline_total_cells += 1
                sig = self._cell_fill_signature(sheet.cell(row=r, column=c))
                if sig:
                    fill_counter[sig] += 1

        background_signatures = set()
        if timeline_total_cells > 0:
            for sig, cnt in fill_counter.items():
                ratio = cnt / float(timeline_total_cells)
                if ratio >= 0.12:
                    background_signatures.add(sig)

        if not background_signatures and fill_counter:
            top_sig, top_cnt = fill_counter.most_common(1)[0]
            top_ratio = top_cnt / float(timeline_total_cells) if timeline_total_cells > 0 else 0.0
            if top_ratio >= 0.35:
                background_signatures.add(top_sig)

        non_bg_signatures = {
            sig for sig, cnt in fill_counter.items()
            if sig not in background_signatures and cnt >= 2
        }

        for r in range(data_start_row, data_end_row + 1):
            task_cells = [self._clean_text(sheet.cell(row=r, column=c).value) for c in left_text_cols]
            task_texts = [t for t in task_cells if t]
            if not task_texts:
                continue
            task_name = task_texts[0]
            task_detail = " / ".join(task_texts[1:4]) if len(task_texts) > 1 else ""

            sig_to_cols: Dict[str, List[int]] = {}
            for c in timeline_cols:
                cell = sheet.cell(row=r, column=c)
                sig = self._cell_fill_signature(cell)
                if not sig:
                    continue
                if sig in non_bg_signatures:
                    sig_to_cols.setdefault(sig, []).append(c)
                elif sig not in background_signatures and self._is_visually_filled(cell):
                    sig_to_cols.setdefault(sig, []).append(c)

            if not sig_to_cols:
                continue

            best_run: List[int] = []
            for _sig, cols in sig_to_cols.items():
                sorted_cols = sorted(cols)
                runs: List[List[int]] = []
                current_run: List[int] = []
                for col in sorted_cols:
                    if not current_run or col == current_run[-1] + 1:
                        current_run.append(col)
                    else:
                        runs.append(current_run)
                        current_run = [col]
                if current_run:
                    runs.append(current_run)
                if not runs:
                    continue
                longest_of_sig = max(runs, key=len)
                if len(longest_of_sig) > len(best_run):
                    best_run = longest_of_sig

            if not best_run:
                continue

            longest = best_run
            if len(longest) < 2:
                continue
            start_col = longest[0]
            end_col = longest[-1]

            confidence = min(
                0.98,
                0.40
                + (0.20 if year_row is not None else 0.0)
                + (0.18 if len(longest) >= 3 else 0.0)
                + (0.12 if len(task_name) >= 2 else 0.0)
                + (0.08 if task_detail else 0.0),
            )
            result.append(
                {
                    "row_no": r,
                    "task_name": task_name,
                    "task_detail": task_detail,
                    "start_col": start_col,
                    "end_col": end_col,
                    "start_label": col_label.get(start_col, f"C{start_col}"),
                    "end_label": col_label.get(end_col, f"C{end_col}"),
                    "confidence": confidence,
                }
            )

        logger.info(
            "XLSX schedule抽取: sheet=%s, timeline_cols=%s, fill_kinds=%s, bg_kinds=%s, rows=%s",
            getattr(sheet, "title", "-"),
            len(timeline_cols),
            len(fill_counter),
            len(background_signatures),
            len(result),
        )
        return result

    def _extract_xlsx_images(
        self,
        file_data: bytes,
        file_name: str,
        file_md5: str,
        user_id: int,
    ) -> List[Dict[str, Any]]:
        """
        抽取 xlsx 内嵌图片并上传 MinIO，同时生成 image block 元数据。
        说明：
        - 仅做图片抽取与邻近单元格描述，不做 VLM。
        - 失败时降级，不影响主文档链路。
        """
        image_blocks: List[Dict[str, Any]] = []
        try:
            wb = load_workbook(BytesIO(file_data), data_only=True, read_only=False)
        except Exception as e:
            logger.warning(f"XLSX 图片抽取：工作簿打开失败: file={file_name}, err={e}")
            return image_blocks

        image_index = 0
        for ws in wb.worksheets:
            ws_images = getattr(ws, "_images", []) or []
            if not ws_images:
                continue
            for img in ws_images:
                try:
                    image_index += 1
                    img_bytes = img._data()  # Raw image bytes from openpyxl
                    if not img_bytes:
                        continue

                    fmt = (getattr(img, "format", None) or "png").lower()
                    if fmt == "jpg":
                        fmt = "jpeg"
                    content_type = f"image/{fmt}"
                    img_md5 = hashlib.md5(img_bytes).hexdigest()

                    anchor_row = None
                    anchor_col = None
                    try:
                        frm = getattr(img.anchor, "_from", None)
                        if frm is not None:
                            anchor_row = int(frm.row) + 1
                            anchor_col = int(frm.col) + 1
                    except Exception:
                        pass

                    description = ""
                    if anchor_row:
                        try:
                            row_values = list(ws.iter_rows(min_row=anchor_row, max_row=anchor_row, values_only=True))
                            if row_values and row_values[0]:
                                texts = [self._clean_text(v) for v in row_values[0] if self._clean_text(v)]
                                description = " | ".join(texts[:6])[:500]
                        except Exception:
                            description = ""
                    if not description:
                        description = f"{ws.title} 图片#{image_index}"

                    object_path = minio_client.build_document_image_path(
                        user_id=user_id,
                        file_md5=file_md5,
                        sheet_name=ws.title,
                        image_index=image_index,
                        ext=fmt,
                    )
                    uploaded = minio_client.upload_bytes(
                        bucket_name=settings.MINIO_DEFAULT_BUCKET,
                        object_name=object_path,
                        data=img_bytes,
                        content_type=content_type,
                    )
                    if not uploaded:
                        logger.warning(f"XLSX 图片上传失败: {file_name}, sheet={ws.title}, idx={image_index}")
                        continue

                    image_blocks.append(
                        {
                            "sheet": ws.title,
                            "anchor_row": anchor_row,
                            "anchor_col": anchor_col,
                            "storage_path": object_path,
                            "image_md5": img_md5,
                            "image_width": int(getattr(img, "width", 0) or 0) or None,
                            "image_height": int(getattr(img, "height", 0) or 0) or None,
                            "content_type": content_type,
                            "description": description,
                            "source_parser": "xlsx_image",
                        }
                    )
                except Exception as img_err:
                    logger.warning(f"XLSX 单图抽取失败: file={file_name}, sheet={ws.title}, err={img_err}")
                    continue

        return image_blocks

    @staticmethod
    def _build_vlm_diagram_prompt(file_name: str, image_name: str) -> str:
        return (
            "あなたは日本企業ドキュメント内の図・フロー図・構成図・画面遷移図を文章化するアシスタントです。"
            "入力画像を読み取り、図の内容を機械検索可能な文章形式へ変換してください。"
            "必ず単一の有効なJSONオブジェクトのみを出力してください。"
            "Markdown、コードフェンス、説明文は禁止です。"
            "目的：図を自然言語で正確に説明し、後続のRAG検索に利用できるテキストへ変換すること。"
            "厳守ルール："
            "1) 画像内で読める情報のみ使用。推測禁止。"
            "2) 読み取れない内容は「不明」。"
            "3) ノード名・画面名・システム名は原文日本語を保持。"
            "4) 存在しない関係を作らない。"
            "5) 順序が不明確なら時系列化しない。"
            "6) すべてのキーを必ず出力。"
            "7) confidenceは0.0〜1.0。"
            "8) page_id と page_name は必須。判別不能でも空文字にせず unknown を入れる。"
            "出力スキーマ:\n"
            "{\n"
            '  "page_id": "画面ID/ページID（例: SCR-USER-LIST。なければ unknown_page）",\n'
            '  "page_name": "画面名/ページ名（なければ unknown）",\n'
            '  "diagram_type": "screen_transition|system_architecture|data_flow|table_layout|timeline|org_chart|network|unknown",\n'
            '  "diagram_overview": "図全体を2〜4文で説明",\n'
            '  "components": ["主要な構成要素1", "主要な構成要素2"],\n'
            '  "process_description": "図内に明確な流れがある場合その手順。なければ該当なし",\n'
            '  "relationships": ["AはBへデータを送信する", "CはDに依存する"],\n'
            '  "keywords": ["重要語句1", "重要語句2"],\n'
            '  "evidence_texts": ["画像から直接読める文字列1", "画像から直接読める文字列2"],\n'
            '  "confidence": 0.0\n'
            "}\n"
            f"文書={file_name}, 画像={image_name}"
        )

    @staticmethod
    def _coerce_confidence(value: Any, default: float = 0.55) -> float:
        try:
            v = float(value)
            if v < 0:
                return 0.0
            if v > 1:
                return 1.0
            return v
        except Exception:
            return default

    def _clean_text_list(self, value: Any, limit: int = 40) -> List[str]:
        if not isinstance(value, list):
            return []
        out: List[str] = []
        seen = set()
        for item in value[:limit]:
            text = self._clean_text(item)
            if not text:
                continue
            key = re.sub(r"\s+", "", text.lower())
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(text)
        return out

    def _parse_edges_from_relationship_lines(self, lines: List[str]) -> List[Dict[str, Any]]:
        edges: List[Dict[str, Any]] = []
        seen = set()
        for line in lines[:80]:
            t = self._clean_text(line)
            if not t:
                continue
            m = re.search(r"(.+?)\s*(?:->|→|⇒|⇢)\s*(.+)", t)
            if not m:
                continue
            src = self._clean_text(m.group(1))
            dst = self._clean_text(m.group(2))
            if not src or not dst or src == dst:
                continue
            src_key = re.sub(r"\s+", "", src.lower())
            dst_key = re.sub(r"\s+", "", dst.lower())
            edge_key = (src_key, dst_key)
            if edge_key in seen:
                continue
            seen.add(edge_key)
            edges.append(
                {
                    "src": src,
                    "dst": dst,
                    "relation_type": "连接",
                    "confidence": 0.70,
                }
            )
        return edges

    @staticmethod
    def _build_fallback_page_id(page_name: str, default_page_id: Optional[str] = None) -> str:
        hint = DocumentProcessorService._clean_text(default_page_id)
        if hint:
            return hint
        name = DocumentProcessorService._clean_text(page_name)
        if not name:
            return "unknown_page"
        slug = re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")
        if slug:
            return slug[:80]
        digest = hashlib.md5(name.encode("utf-8")).hexdigest()[:8]
        return f"page_{digest}"

    def _normalize_vlm_diagram_payload(
        self,
        payload: Dict[str, Any],
        default_page_name: Optional[str] = None,
        default_page_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        兼容新旧 VLM 输出：
        - 新：diagram_overview/components/process_description/relationships/keywords/evidence_texts
        - 旧：diagram_summary/nodes/edges
        """
        page_name = self._clean_text(
            payload.get("page_name")
            or payload.get("screen_name")
            or payload.get("view_name")
            or default_page_name
            or "unknown"
        )
        page_id = self._clean_text(
            payload.get("page_id")
            or payload.get("screen_id")
            or payload.get("view_id")
            or default_page_id
        )
        if not page_id:
            page_id = self._build_fallback_page_id(page_name, default_page_id=default_page_id)
        diagram_type = self._clean_text(payload.get("diagram_type")) or "unknown"
        overview = self._clean_text(payload.get("diagram_overview") or payload.get("diagram_summary"))
        process_desc = self._clean_text(payload.get("process_description"))
        components = self._clean_text_list(payload.get("components"), limit=50)
        relationships = self._clean_text_list(payload.get("relationships"), limit=80)
        keywords = self._clean_text_list(payload.get("keywords"), limit=50)
        evidence_texts = self._clean_text_list(payload.get("evidence_texts"), limit=80)
        confidence = self._coerce_confidence(payload.get("confidence"), default=0.65)

        node_names: List[str] = []
        seen_nodes = set()

        for name in components:
            key = re.sub(r"\s+", "", name.lower())
            if not key or key in seen_nodes:
                continue
            seen_nodes.add(key)
            node_names.append(name)

        nodes_raw = payload.get("nodes") if isinstance(payload.get("nodes"), list) else []
        for row in nodes_raw[:80]:
            if not isinstance(row, dict):
                continue
            name = self._clean_text(row.get("name"))
            if not name:
                continue
            key = re.sub(r"\s+", "", name.lower())
            if not key or key in seen_nodes:
                continue
            seen_nodes.add(key)
            node_names.append(name)

        edges: List[Dict[str, Any]] = []
        seen_edges = set()
        edges_raw = payload.get("edges") if isinstance(payload.get("edges"), list) else []
        for row in edges_raw[:120]:
            if not isinstance(row, dict):
                continue
            src = self._clean_text(row.get("src"))
            dst = self._clean_text(row.get("dst"))
            if not src or not dst or src == dst:
                continue
            rel = self._clean_text(row.get("relation_type")) or "连接"
            conf = self._coerce_confidence(row.get("confidence"), default=confidence)
            edge_key = (
                re.sub(r"\s+", "", src.lower()),
                re.sub(r"\s+", "", dst.lower()),
                rel,
            )
            if edge_key in seen_edges:
                continue
            seen_edges.add(edge_key)
            edges.append({"src": src, "dst": dst, "relation_type": rel, "confidence": conf})

        rel_edges = self._parse_edges_from_relationship_lines(relationships)
        for e in rel_edges:
            edge_key = (
                re.sub(r"\s+", "", e["src"].lower()),
                re.sub(r"\s+", "", e["dst"].lower()),
                e["relation_type"],
            )
            if edge_key in seen_edges:
                continue
            seen_edges.add(edge_key)
            edges.append(e)

        summary_lines: List[str] = []
        summary_lines.append(f"ページID: {page_id}")
        summary_lines.append(f"ページ名: {page_name}")
        summary_lines.append(f"図タイプ: {diagram_type}")
        if overview:
            summary_lines.append(overview)
        if components:
            summary_lines.append(f"主要構成要素: {'、'.join(components[:20])}")
        if process_desc and process_desc not in {"該当なし", "なし", "不明"}:
            summary_lines.append(f"処理/流れ: {process_desc}")
        if relationships:
            summary_lines.append(f"関係: {' / '.join(relationships[:20])}")
        if keywords:
            summary_lines.append(f"キーワード: {'、'.join(keywords[:30])}")
        if evidence_texts:
            summary_lines.append(f"図中テキスト: {' / '.join(evidence_texts[:20])}")
        summary_text = self._clean_text("\n".join(summary_lines))

        return {
            "page_id": page_id,
            "page_name": page_name,
            "diagram_type": diagram_type,
            "summary_text": summary_text,
            "nodes": node_names,
            "edges": edges,
            "confidence": confidence,
        }

    def _extract_office_media_images(
        self,
        file_data: bytes,
        ext: str,
        max_images: int = 4,
    ) -> List[Dict[str, Any]]:
        """
        从 office 压缩包提取媒体图片字节（用于 VLM 结构化）。
        """
        ext_norm = (ext or "").lower()
        if ext_norm == "docx":
            prefix = "word/media/"
        elif ext_norm in ("xlsx", "xlsm", "xltx", "xltm"):
            prefix = "xl/media/"
        else:
            return []

        try:
            zf = zipfile.ZipFile(BytesIO(file_data))
        except Exception:
            return []

        image_names = sorted(
            [
                n for n in zf.namelist()
                if n.startswith(prefix) and re.search(r"\.(png|jpe?g|bmp|webp)$", n, re.IGNORECASE)
            ]
        )[:max_images]
        images: List[Dict[str, Any]] = []
        for name in image_names:
            try:
                raw = zf.read(name)
                if not raw:
                    continue
                img = Image.open(BytesIO(raw))
                fmt = (img.format or "PNG").upper()
                if fmt not in {"PNG", "JPEG", "JPG", "WEBP", "BMP"}:
                    continue
                w, h = img.size
                if (w * h) < 120000 or min(w, h) < 220:
                    continue
                max_side = 1800
                if max(w, h) > max_side:
                    ratio = max_side / float(max(w, h))
                    img = img.resize((max(1, int(w * ratio)), max(1, int(h * ratio))))
                buf = BytesIO()
                out_fmt = "JPEG" if fmt in {"JPG", "JPEG"} else "PNG"
                img.convert("RGB").save(buf, format=out_fmt, quality=92)
                mime = "image/jpeg" if out_fmt == "JPEG" else "image/png"
                images.append(
                    {
                        "name": Path(name).name,
                        "bytes": buf.getvalue(),
                        "mime_type": mime,
                    }
                )
            except Exception:
                continue
        return images

    async def _extract_office_vlm_diagram_blocks(
        self,
        file_data: bytes,
        file_name: str,
        file_md5: str,
        user_id: int,
        ext: str,
        start_index: int,
        max_images: int = 3,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        对 Office 文档内嵌图片执行 VLM 结构化抽取，产出 diagram_* blocks。
        """
        if not settings.OPENAI_VISION_ENABLED:
            return [], []
        if not settings.OPENAI_API_KEY:
            return [], []

        images = self._extract_office_media_images(file_data=file_data, ext=ext, max_images=max_images)
        if not images:
            return [], []

        blocks: List[Dict[str, Any]] = []
        image_blocks: List[Dict[str, Any]] = []
        block_idx = start_index
        for img_idx, img in enumerate(images, 1):
            try:
                image_ext = "jpg" if str(img.get("mime_type") or "").lower().endswith("jpeg") else "png"
                object_path = minio_client.build_document_image_path(
                    user_id=user_id,
                    file_md5=file_md5,
                    sheet_name="vlm_office",
                    image_index=9000 + img_idx,
                    ext=image_ext,
                )
                uploaded = minio_client.upload_bytes(
                    bucket_name=settings.MINIO_DEFAULT_BUCKET,
                    object_name=object_path,
                    data=img["bytes"],
                    content_type=img.get("mime_type") or "image/png",
                )
                if not uploaded:
                    logger.warning(
                        "VLM_OFFICE_IMAGE_UPLOAD_FAIL: file=%s, file_md5=%s, image=%s",
                        file_name,
                        file_md5,
                        img.get("name"),
                    )
                    continue
                img_w = 0
                img_h = 0
                try:
                    with Image.open(BytesIO(img["bytes"])) as im:
                        img_w, img_h = im.size
                except Exception:
                    pass
                image_blocks.append(
                    {
                        "sheet": None,
                        "page": None,
                        "anchor_row": None,
                        "anchor_col": None,
                        "storage_path": object_path,
                        "image_md5": hashlib.md5(img["bytes"]).hexdigest(),
                        "image_width": img_w or None,
                        "image_height": img_h or None,
                        "content_type": img.get("mime_type") or "image/png",
                        "description": f"office vlm source image: {img.get('name') or 'embedded_image'}",
                        "source_parser": "vlm_diagram",
                    }
                )
                prompt = self._build_vlm_diagram_prompt(file_name=file_name, image_name=img["name"])
                payload = await self.chat_client.analyze_image_json(
                    image_bytes=img["bytes"],
                    mime_type=img["mime_type"],
                    prompt=prompt,
                )
                if not payload:
                    continue
                logger.info(
                    "VLM_RAW_PAYLOAD_JSON: file=%s, image=%s, scope=office_media, payload=%s",
                    file_name,
                    img.get("name"),
                    json.dumps(payload, ensure_ascii=False),
                )

                normalized = self._normalize_vlm_diagram_payload(
                    payload,
                    default_page_name=img.get("name"),
                )
                logger.info(
                    "VLM_OVERVIEW_FULL: file=%s, image=%s, scope=office_media, diagram_overview=%s, process_description=%s, relationships=%s, evidence_texts=%s",
                    file_name,
                    img.get("name"),
                    self._clean_text(payload.get("diagram_overview") or payload.get("diagram_summary")),
                    self._clean_text(payload.get("process_description")),
                    json.dumps(self._clean_text_list(payload.get("relationships"), limit=80), ensure_ascii=False),
                    json.dumps(self._clean_text_list(payload.get("evidence_texts"), limit=80), ensure_ascii=False),
                )
                logger.info(
                    "VLM_NORMALIZED_FULL_SUMMARY: file=%s, image=%s, scope=office_media, summary_text=%s",
                    file_name,
                    img.get("name"),
                    self._clean_text(normalized.get("summary_text")),
                )
                page_id = self._clean_text(normalized.get("page_id")) or "unknown_page"
                page_name = self._clean_text(normalized.get("page_name")) or self._clean_text(img.get("name")) or "unknown"
                diagram_type = self._clean_text(normalized.get("diagram_type")) or "unknown"
                summary = self._clean_text(normalized.get("summary_text"))
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "diagram_page",
                        "text": (
                            f"[diagram_page] page_id={page_id} ; page_name={page_name} ; "
                            f"diagram_type={diagram_type}"
                        ),
                        "page": None,
                        "section": page_name or img["name"],
                        "sheet": None,
                        "source_parser": "vlm_diagram",
                        "file_type": ext,
                        "file_name": file_name,
                        "image_path": object_path,
                    }
                )
                block_idx += 1
                if summary:
                    blocks.append(
                        {
                            "block_index": block_idx,
                            "type": "diagram_summary",
                            "text": f"[diagram_summary] {summary}",
                            "page": None,
                            "section": page_name or img["name"],
                            "sheet": None,
                            "source_parser": "vlm_diagram",
                            "file_type": ext,
                            "file_name": file_name,
                            "image_path": object_path,
                        }
                    )
                    block_idx += 1

                node_names: List[str] = []
                seen_nodes = set()
                for name in normalized.get("nodes", [])[:80]:
                    name = self._clean_text(name)
                    if not name:
                        continue
                    key = re.sub(r"\s+", "", name.lower())
                    if not key or key in seen_nodes:
                        continue
                    seen_nodes.add(key)
                    node_names.append(name)
                    blocks.append(
                        {
                            "block_index": block_idx,
                            "type": "diagram_node",
                            "text": f"[diagram_node] {name}",
                            "page": None,
                            "section": page_name or img["name"],
                            "sheet": None,
                            "source_parser": "vlm_diagram",
                            "file_type": ext,
                            "file_name": file_name,
                            "image_path": object_path,
                        }
                    )
                    block_idx += 1

                if node_names:
                    node_set_norm = {re.sub(r"\s+", "", n.lower()) for n in node_names}
                else:
                    node_set_norm = set()

                seen_edges = set()
                for row in normalized.get("edges", [])[:120]:
                    src = self._clean_text(row.get("src"))
                    dst = self._clean_text(row.get("dst"))
                    if not src or not dst or src == dst:
                        continue
                    src_key = re.sub(r"\s+", "", src.lower())
                    dst_key = re.sub(r"\s+", "", dst.lower())
                    if node_set_norm and (src_key not in node_set_norm or dst_key not in node_set_norm):
                        for miss in (src, dst):
                            mk = re.sub(r"\s+", "", miss.lower())
                            if mk and mk not in node_set_norm:
                                node_set_norm.add(mk)
                                blocks.append(
                                    {
                                        "block_index": block_idx,
                                        "type": "diagram_node",
                                        "text": f"[diagram_node] {miss}",
                                        "page": None,
                                        "section": page_name or img["name"],
                                        "sheet": None,
                                        "source_parser": "vlm_diagram",
                                        "file_type": ext,
                                        "file_name": file_name,
                                        "image_path": object_path,
                                    }
                                )
                                block_idx += 1
                    rel = self._clean_text(row.get("relation_type")) or "连接"
                    conf = self._coerce_confidence(row.get("confidence"), default=normalized.get("confidence", 0.65))
                    edge_key = (src_key, dst_key, rel)
                    if edge_key in seen_edges:
                        continue
                    seen_edges.add(edge_key)
                    blocks.append(
                        {
                            "block_index": block_idx,
                            "type": "diagram_edge",
                            "text": f"{src} -> {dst} ; rel={rel} ; confidence={conf:.2f}",
                            "page": None,
                            "section": page_name or img["name"],
                            "sheet": None,
                            "source_parser": "vlm_diagram",
                            "file_type": ext,
                            "file_name": file_name,
                            "image_path": object_path,
                        }
                    )
                    block_idx += 1
            except Exception as e:
                logger.warning("VLM图形解析失败: file=%s, image=%s, err=%s", file_name, img.get("name"), e)
                continue

        if blocks:
            logger.info("VLM图形解析完成: file=%s, image_count=%s, blocks=%s", file_name, len(images), len(blocks))
        return blocks, image_blocks

    @staticmethod
    def _load_diagram_font(size: int = 24) -> Any:
        candidates = [
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for path in candidates:
            try:
                return ImageFont.truetype(path, size=size)
            except Exception:
                continue
        return ImageFont.load_default()

    def _collect_xlsx_diagram_graphs_from_blocks(self, blocks: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
        graphs: Dict[str, Dict[str, Any]] = {}
        for b in blocks or []:
            if str(b.get("source_parser") or "") != "xlsx_diagram":
                continue
            sheet = str(b.get("sheet") or "sheet")
            g = graphs.setdefault(sheet, {"nodes": [], "edges": []})
            btype = str(b.get("type") or "")
            if btype == "diagram_node":
                row_json = b.get("row_json") if isinstance(b.get("row_json"), dict) else {}
                name = self._clean_text(
                    row_json.get("name")
                    or str(b.get("text") or "").replace("[diagram_node]", "", 1)
                )
                if not name:
                    continue
                g["nodes"].append(
                    {
                        "name": name,
                        "center_x": row_json.get("center_x"),
                        "center_y": row_json.get("center_y"),
                        "left": row_json.get("left"),
                        "top": row_json.get("top"),
                        "right": row_json.get("right"),
                        "bottom": row_json.get("bottom"),
                    }
                )
            elif btype == "diagram_edge":
                row_json = b.get("row_json") if isinstance(b.get("row_json"), dict) else {}
                src = self._clean_text(row_json.get("src"))
                dst = self._clean_text(row_json.get("dst"))
                rel = self._clean_text(row_json.get("relation_type")) or "连接"
                if not src or not dst:
                    txt = str(b.get("text") or "")
                    m = re.search(r"(.+?)\s*->\s*(.+?)(?:\s*;|$)", txt)
                    if m:
                        src = self._clean_text(m.group(1))
                        dst = self._clean_text(m.group(2))
                if src and dst:
                    g["edges"].append({"src": src, "dst": dst, "relation_type": rel})

        for sheet, g in graphs.items():
            uniq = {}
            for n in g["nodes"]:
                k = re.sub(r"\s+", "", str(n.get("name") or "").lower())
                if k and k not in uniq:
                    uniq[k] = n
            g["nodes"] = list(uniq.values())
        return graphs

    def _render_xlsx_diagram_graph_image(self, sheet: str, graph: Dict[str, Any]) -> Optional[bytes]:
        nodes = graph.get("nodes") or []
        if len(nodes) < 3:
            return None
        edges = graph.get("edges") or []

        canvas_w, canvas_h = 1800, 1200
        margin_x, margin_y = 80, 80
        img = Image.new("RGB", (canvas_w, canvas_h), "white")
        draw = ImageDraw.Draw(img)
        font_title = self._load_diagram_font(36)
        font_node = self._load_diagram_font(22)
        font_edge = self._load_diagram_font(18)
        draw.text((margin_x, 20), f"Sheet: {sheet}", fill="#111827", font=font_title)

        def _to_float(v: Any, default: float = 0.0) -> float:
            try:
                return float(v)
            except Exception:
                return default

        xs = [_to_float(n.get("center_x"), 0.0) for n in nodes if n.get("center_x") is not None]
        ys = [_to_float(n.get("center_y"), 0.0) for n in nodes if n.get("center_y") is not None]
        use_geo = bool(xs and ys)
        node_pos: Dict[str, Tuple[float, float]] = {}
        if use_geo:
            min_x, max_x = min(xs), max(xs)
            min_y, max_y = min(ys), max(ys)
            span_x = max(1.0, max_x - min_x)
            span_y = max(1.0, max_y - min_y)
            for i, n in enumerate(nodes):
                name = str(n.get("name") or f"node_{i}")
                cx = _to_float(n.get("center_x"), min_x)
                cy = _to_float(n.get("center_y"), min_y)
                px = margin_x + ((cx - min_x) / span_x) * (canvas_w - margin_x * 2)
                py = margin_y + 60 + ((cy - min_y) / span_y) * (canvas_h - margin_y * 2 - 60)
                node_pos[name] = (px, py)
        else:
            cols = 4
            col_w = (canvas_w - margin_x * 2) / cols
            row_h = 140
            for i, n in enumerate(nodes):
                name = str(n.get("name") or f"node_{i}")
                c = i % cols
                r = i // cols
                node_pos[name] = (margin_x + col_w * c + col_w / 2, margin_y + 100 + row_h * r)

        if not edges:
            buckets: Dict[int, List[Tuple[str, Tuple[float, float]]]] = {}
            for n, pos in node_pos.items():
                k = int(pos[1] // 140)
                buckets.setdefault(k, []).append((n, pos))
            for _, arr in buckets.items():
                arr.sort(key=lambda x: x[1][0])
                for i in range(len(arr) - 1):
                    edges.append({"src": arr[i][0], "dst": arr[i + 1][0], "relation_type": "连接"})

        for e in edges[:120]:
            src = str(e.get("src") or "")
            dst = str(e.get("dst") or "")
            if src not in node_pos or dst not in node_pos or src == dst:
                continue
            x1, y1 = node_pos[src]
            x2, y2 = node_pos[dst]
            draw.line((x1, y1, x2, y2), fill="#2563eb", width=4)
            # arrow head
            vx, vy = (x2 - x1), (y2 - y1)
            norm = max((vx * vx + vy * vy) ** 0.5, 1.0)
            ux, uy = vx / norm, vy / norm
            ah = 14
            px, py = -uy, ux
            p1 = (x2, y2)
            p2 = (x2 - ux * ah + px * ah * 0.5, y2 - uy * ah + py * ah * 0.5)
            p3 = (x2 - ux * ah - px * ah * 0.5, y2 - uy * ah - py * ah * 0.5)
            draw.polygon([p1, p2, p3], fill="#2563eb")
            rel = str(e.get("relation_type") or "").strip()
            if rel:
                mx, my = (x1 + x2) / 2, (y1 + y2) / 2
                draw.text((mx + 4, my + 4), rel[:10], fill="#374151", font=font_edge)

        for name, (cx, cy) in node_pos.items():
            w = 240
            h = 72
            x1, y1 = cx - w / 2, cy - h / 2
            x2, y2 = cx + w / 2, cy + h / 2
            draw.rounded_rectangle((x1, y1, x2, y2), radius=10, outline="#f59e0b", width=3, fill="#fff7ed")
            draw.text((x1 + 10, y1 + 20), name[:24], fill="#111827", font=font_node)

        out = BytesIO()
        img.save(out, format="PNG")
        return out.getvalue()

    def _render_office_pages_via_libreoffice(
        self,
        file_data: bytes,
        file_name: str,
        max_pages: int = 4,
    ) -> List[Dict[str, Any]]:
        """
        使用 LibreOffice 进行高保真渲染：
        office -> pdf -> png(page)。
        这是“整页快照”，优先于重绘图用于 VLM。
        """
        if not fitz:
            return []

        ext = Path(file_name).suffix.lower().lstrip(".")
        if ext not in ("xlsx", "xlsm", "xltx", "xltm", "docx", "pptx"):
            return []

        images: List[Dict[str, Any]] = []
        try:
            with tempfile.TemporaryDirectory(prefix="akb_office_render_") as td:
                input_path = Path(td) / file_name
                input_path.write_bytes(file_data)

                cmd = [
                    "soffice",
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--nolockcheck",
                    "--norestore",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    td,
                    str(input_path),
                ]
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=90)
                if proc.returncode != 0:
                    logger.warning("LibreOffice 转换失败: file=%s, stderr=%s", file_name, (proc.stderr or "").strip())
                    return []

                pdf_path = Path(td) / f"{Path(file_name).stem}.pdf"
                if not pdf_path.exists():
                    candidates = sorted(Path(td).glob("*.pdf"))
                    if not candidates:
                        return []
                    pdf_path = candidates[0]

                doc = fitz.open(str(pdf_path))
                try:
                    total = min(max_pages, doc.page_count)
                    for i in range(total):
                        page = doc.load_page(i)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                        png_bytes = pix.tobytes("png")
                        images.append(
                            {
                                "name": f"page_{i+1}.png",
                                "bytes": png_bytes,
                                "mime_type": "image/png",
                                "page": i + 1,
                            }
                        )
                finally:
                    doc.close()
        except Exception as e:
            logger.warning("LibreOffice 渲染异常: file=%s, err=%s", file_name, e)
            return []

        return images

    def _render_xlsx_sheets_via_libreoffice(
        self,
        file_data: bytes,
        file_name: str,
        target_sheets: List[str],
        blocks: Optional[List[Dict[str, Any]]] = None,
        max_pages_per_sheet: int = 3,
    ) -> List[Dict[str, Any]]:
        """
        按 sheet 维度渲染 XLSX 快照：
        - 直接对原始工作簿做 LibreOffice 渲染（不经 openpyxl 重写，避免图层丢失）
        - 从导出 PDF 页文本映射 page -> sheet
        - 仅保留命中目标 sheet 的页面
        """
        if not fitz:
            return []
        ext = Path(file_name).suffix.lower().lstrip(".")
        if ext not in ("xlsx", "xlsm", "xltx", "xltm"):
            return []
        target_sheet_set = {str(s).strip() for s in (target_sheets or []) if str(s).strip()}
        if not target_sheet_set:
            return []

        sheet_tokens = self._build_sheet_match_tokens(blocks or [], target_sheet_set)
        logger.info(
            "XLSX snapshot render start: file=%s, target_sheets=%s, token_sheets=%s",
            file_name,
            sorted(list(target_sheet_set)),
            {k: len(v) for k, v in sheet_tokens.items()},
        )
        snapshots: List[Dict[str, Any]] = []
        sheet_page_counter: Dict[str, int] = {}
        page_rows: List[Dict[str, Any]] = []

        try:
            with tempfile.TemporaryDirectory(prefix="akb_xlsx_sheet_render_") as td:
                xlsx_path = Path(td) / file_name
                xlsx_path.write_bytes(file_data)

                cmd = [
                    "soffice",
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--nolockcheck",
                    "--norestore",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    td,
                    str(xlsx_path),
                ]
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
                if proc.returncode != 0:
                    logger.warning(
                        "XLSX 原始工作簿快照转换失败: file=%s, stderr=%s",
                        file_name,
                        (proc.stderr or "").strip(),
                    )
                    return []

                pdf_files = sorted(Path(td).glob("*.pdf"))
                if not pdf_files:
                    return []
                pdf_path = pdf_files[0]

                doc = fitz.open(str(pdf_path))
                try:
                    logger.info(
                        "XLSX snapshot pdf pages: file=%s, pages=%s",
                        file_name,
                        doc.page_count,
                    )
                    for i in range(doc.page_count):
                        page = doc.load_page(i)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                        png_bytes = pix.tobytes("png")

                        page_text = self._clean_text(page.get_text("text"))
                        matched_sheet = self._match_sheet_by_name(page_text, sheet_tokens)
                        matched_mode = "text_strict"
                        match_confidence = 1.00
                        if not matched_sheet:
                            ocr_text = self._ocr_image_text(png_bytes)
                            matched_sheet = self._match_sheet_by_name(ocr_text, sheet_tokens)
                            matched_mode = "ocr_strict" if matched_sheet else "none"
                            if matched_sheet:
                                match_confidence = 0.95
                        else:
                            ocr_text = ""

                        if not matched_sheet and len(target_sheet_set) == 1:
                            matched_sheet = next(iter(target_sheet_set))
                            matched_mode = "single_candidate"
                            match_confidence = 0.35

                        match_reason = ""
                        if matched_mode == "text_strict":
                            match_reason = "sheet name matched in page text"
                        elif matched_mode == "ocr_strict":
                            match_reason = "sheet name matched in OCR text"
                        elif matched_mode == "single_candidate":
                            match_reason = "single candidate fallback"
                        else:
                            match_reason = "no match"

                        page_rows.append(
                            {
                                "page": i + 1,
                                "png_bytes": png_bytes,
                                "sheet": matched_sheet,
                                "mode": matched_mode,
                                "confidence": match_confidence if matched_sheet else 0.0,
                                "reason": match_reason,
                            }
                        )
                        logger.info(
                            "XLSX snapshot page match: file=%s, page=%s, text_len=%s, matched_sheet=%s, mode=%s, confidence=%.2f",
                            file_name,
                            i + 1,
                            len(page_text or ""),
                            matched_sheet or "-",
                            matched_mode,
                            match_confidence,
                        )
                finally:
                    doc.close()
        except Exception as e:
            logger.warning("XLSX sheet快照异常: file=%s, err=%s", file_name, e)
            return []

        for idx, row in enumerate(page_rows):
            if row.get("sheet"):
                continue
            prev_sheet = page_rows[idx - 1].get("sheet") if idx > 0 else None
            next_sheet = page_rows[idx + 1].get("sheet") if idx + 1 < len(page_rows) else None
            if prev_sheet and prev_sheet == next_sheet and prev_sheet in target_sheet_set:
                row["sheet"] = prev_sheet
                row["mode"] = "neighbor_window"
                row["confidence"] = 0.45
                row["reason"] = "neighbor window fallback"
                logger.info(
                    "XLSX snapshot page backfill(neighbor_window): file=%s, page=%s, sheet=%s",
                    file_name,
                    row.get("page"),
                    prev_sheet,
                )
            else:
                row["mode"] = "unbound"
                row["confidence"] = 0.0
                row["reason"] = "unbound after all matching strategies"

        accepted_modes = {"text_strict", "ocr_strict", "neighbor_window", "single_candidate"}
        skipped_unbound = 0
        skipped_low_conf = 0
        for row in page_rows:
            matched_sheet = str(row.get("sheet") or "").strip()
            if matched_sheet and matched_sheet not in target_sheet_set:
                logger.info(
                    "XLSX snapshot page skipped(not target): file=%s, page=%s, matched_sheet=%s",
                    file_name,
                    row.get("page"),
                    matched_sheet,
                )
                continue
            mode = str(row.get("mode") or "").strip().lower()
            if not matched_sheet:
                skipped_unbound += 1
                logger.info(
                    "XLSX snapshot page skipped(unbound): file=%s, page=%s, mode=%s, reason=%s",
                    file_name,
                    row.get("page"),
                    mode or "-",
                    row.get("reason") or "-",
                )
                continue
            if mode not in accepted_modes:
                skipped_low_conf += 1
                logger.info(
                    "XLSX snapshot page skipped(low_conf_mode): file=%s, page=%s, sheet=%s, mode=%s, confidence=%s",
                    file_name,
                    row.get("page"),
                    matched_sheet,
                    mode or "-",
                    row.get("confidence"),
                )
                continue

            counter_key = matched_sheet if matched_sheet else "__unbound__"
            count = sheet_page_counter.get(counter_key, 0)
            if count >= max_pages_per_sheet:
                continue
            sheet_page_counter[counter_key] = count + 1

            page_no = int(row.get("page") or 0) or 1
            safe_sheet_name = (
                re.sub(r'[\\\\/:*?"<>|]+', "_", matched_sheet).strip() if matched_sheet else "unbound"
            ) or "sheet"
            snapshots.append(
                {
                    "name": f"{safe_sheet_name}_page_{page_no}.png",
                    "bytes": row.get("png_bytes") or b"",
                    "mime_type": "image/png",
                    "sheet": matched_sheet or None,
                    "page": page_no,
                    "match_mode": row.get("mode"),
                    "match_confidence": row.get("confidence"),
                    "match_reason": row.get("reason"),
                }
            )

        logger.info(
            "XLSX 原始工作簿快照完成: file=%s, target_sheets=%s, snapshots=%s, mapped_sheets=%s, skipped_unbound=%s, skipped_low_conf=%s",
            file_name,
            len(target_sheet_set),
            len(snapshots),
            len(sheet_page_counter),
            skipped_unbound,
            skipped_low_conf,
        )
        return snapshots

    @staticmethod
    def _normalize_text_for_match(text: str) -> str:
        t = DocumentProcessorService._clean_text(text or "")
        t = t.lower()
        t = t.replace("　", " ").replace("・", "").replace(" ", "")
        t = t.replace("-", "").replace("_", "")
        return t

    def _ocr_image_text(self, image_bytes: bytes) -> str:
        if pytesseract is None:
            return ""
        try:
            with Image.open(BytesIO(image_bytes)) as im:
                text = pytesseract.image_to_string(im, lang="jpn+eng")
                return self._clean_text(text)
        except Exception:
            return ""

    def _collect_snapshot_candidate_sheets(self, blocks: List[Dict[str, Any]]) -> set[str]:
        """
        为 xlsx 快照挑选“真正图形化”的 sheet。
        规则（避免普通表格 sheet 被误纳入）：
        - 至少 1 条 diagram_edge 且至少 2 个 diagram_node；或
        - 至少 settings.XLSX_SNAPSHOT_NODE_THRESHOLD 个 diagram_node（允许无 summary）。
        - 低文本密度兜底（疑似图形页）：
          * sheet 非空
          * text_chars / non_empty_cells 较低
          * distinct_text_tokens 不高
        """
        stat: Dict[str, Dict[str, Any]] = {}
        for b in blocks or []:
            sheet = str(b.get("sheet") or "").strip()
            if not sheet:
                continue
            btype = str(b.get("type") or "")
            s = stat.setdefault(
                sheet,
                {
                    "edge": 0,
                    "node": 0,
                    "summary": 0,
                    "table_rows": 0,
                    "text_chars": 0,
                    "token_count": 0,
                    "non_empty_cells": 0,
                    "distinct_tokens": set(),
                },
            )
            text = self._clean_text(b.get("text"))
            if text:
                s["text_chars"] += len(text)
                tokens = re.findall(r"[\u3040-\u30ff\u3400-\u9fffA-Za-z0-9]{2,24}", text)
                token_count = len(tokens)
                s["token_count"] += token_count
                s["distinct_tokens"].update(t.lower() for t in tokens)
                s["non_empty_cells"] += 1

            row_json = b.get("row_json")
            if isinstance(row_json, dict) and row_json:
                non_empty = sum(1 for v in row_json.values() if self._clean_text(v))
                if non_empty > 0:
                    s["non_empty_cells"] += non_empty
            if btype in {"table_row", "table_header", "schedule_row"}:
                s["table_rows"] += 1
            if btype == "diagram_edge":
                s["edge"] += 1
            elif btype == "diagram_node":
                s["node"] += 1
            elif btype == "diagram_summary":
                s["summary"] += 1

        picked: set[str] = set()
        node_threshold = max(1, int(settings.XLSX_SNAPSHOT_NODE_THRESHOLD))
        low_chars_max = max(1, int(settings.XLSX_SNAPSHOT_LOW_TEXT_CHARS_MAX))
        low_token_max = max(1, int(settings.XLSX_SNAPSHOT_LOW_TOKEN_MAX))
        low_char_per_non_empty = max(0.1, float(settings.XLSX_SNAPSHOT_LOW_CHAR_PER_NON_EMPTY_MAX))
        low_distinct_token_max = max(1, int(settings.XLSX_SNAPSHOT_LOW_DISTINCT_TOKEN_MAX))
        low_non_empty_min = max(1, int(settings.XLSX_SNAPSHOT_LOW_NON_EMPTY_MIN))
        for sheet, s in stat.items():
            distinct_token_count = len(s.get("distinct_tokens") or set())
            non_empty_cells = int(s.get("non_empty_cells") or 0)
            avg_chars_per_non_empty = (
                float(s["text_chars"]) / float(non_empty_cells)
                if non_empty_cells > 0
                else 9999.0
            )
            low_text_density = (
                non_empty_cells >= low_non_empty_min
                and s["text_chars"] <= low_chars_max
                and s["token_count"] <= low_token_max
                and avg_chars_per_non_empty <= low_char_per_non_empty
                and distinct_token_count <= low_distinct_token_max
                and s["table_rows"] <= 2
            )
            has_graph_pattern = (s["edge"] >= 1 and s["node"] >= 2) or s["node"] >= node_threshold
            if has_graph_pattern or low_text_density:
                picked.add(sheet)
            s["distinct_token_count"] = distinct_token_count
            s["avg_chars_per_non_empty"] = round(avg_chars_per_non_empty, 2)
            if "distinct_tokens" in s:
                s.pop("distinct_tokens", None)
        logger.info(
            "XLSX snapshot candidate sheets: total_sheets=%s, picked=%s, detail=%s",
            len(stat),
            sorted(list(picked)),
            stat,
        )
        return picked

    def _match_sheet_by_name(
        self,
        snapshot_text: str,
        sheet_tokens: Dict[str, set[str]],
    ) -> Optional[str]:
        """
        严格匹配：必须命中 sheet 名本身。
        """
        if not snapshot_text or not sheet_tokens:
            return None
        ntext = self._normalize_text_for_match(snapshot_text)
        if not ntext:
            return None

        strict_candidates: List[str] = []
        for sheet in sheet_tokens.keys():
            ns = self._normalize_text_for_match(sheet)
            if ns and ns in ntext:
                strict_candidates.append(sheet)
        if len(strict_candidates) == 1:
            return strict_candidates[0]
        if len(strict_candidates) > 1:
            best_sheet = None
            best_score = -1
            for sheet in strict_candidates:
                toks = sheet_tokens.get(sheet) or set()
                score = 0
                for tok in list(toks)[:120]:
                    if tok and tok in ntext:
                        score += 1
                if score > best_score:
                    best_score = score
                    best_sheet = sheet
            return best_sheet
        return None

    def _match_snapshot_sheet_weak(
        self,
        snapshot_text: str,
        sheet_tokens: Dict[str, set[str]],
    ) -> Optional[str]:
        """
        弱匹配：不要求 sheet 名直命中，按 token 打分并要求明显领先。
        """
        if not snapshot_text or not sheet_tokens:
            return None
        ntext = self._normalize_text_for_match(snapshot_text)
        if not ntext:
            return None
        scored: List[Tuple[str, int]] = []
        for sheet, toks in sheet_tokens.items():
            score = 0
            for tok in list(toks)[:180]:
                if tok and tok in ntext:
                    score += 1
            if score > 0:
                scored.append((sheet, score))
        if not scored:
            return None
        scored.sort(key=lambda x: x[1], reverse=True)
        best_sheet, best_score = scored[0]
        second_score = scored[1][1] if len(scored) > 1 else -1
        min_score = max(1, int(settings.XLSX_SNAPSHOT_WEAK_MATCH_MIN_SCORE))
        min_lead = max(0, int(settings.XLSX_SNAPSHOT_WEAK_MATCH_MIN_LEAD))
        if best_score >= min_score and best_score >= second_score + min_lead:
            return best_sheet
        return None

    def _build_sheet_match_tokens(
        self,
        blocks: List[Dict[str, Any]],
        candidate_sheets: set[str],
    ) -> Dict[str, set[str]]:
        tokens: Dict[str, set[str]] = {}
        for sheet in candidate_sheets:
            normalized_name = self._normalize_text_for_match(sheet)
            s = tokens.setdefault(sheet, set())
            if normalized_name:
                s.add(normalized_name)

        for b in blocks or []:
            sheet = str(b.get("sheet") or "").strip()
            if sheet not in candidate_sheets:
                continue
            text = self._clean_text(b.get("text"))
            if not text:
                continue
            words = re.findall(r"[\u3040-\u30ff\u3400-\u9fffA-Za-z0-9]{2,16}", text)
            for w in words[:40]:
                nw = self._normalize_text_for_match(w)
                if nw and len(nw) >= 2:
                    tokens[sheet].add(nw)
        return tokens

    def _match_snapshot_sheet(
        self,
        snapshot_ocr_text: str,
        sheet_tokens: Dict[str, set[str]],
    ) -> Optional[str]:
        if not snapshot_ocr_text or not sheet_tokens:
            return None
        ntext = self._normalize_text_for_match(snapshot_ocr_text)
        if not ntext:
            return None
        strict_candidates: List[str] = []
        for sheet in sheet_tokens.keys():
            ns = self._normalize_text_for_match(sheet)
            if ns and ns in ntext:
                strict_candidates.append(sheet)
        if len(strict_candidates) == 1:
            return strict_candidates[0]
        if len(strict_candidates) > 1:
            best_sheet = None
            best_score = -1
            for sheet in strict_candidates:
                toks = sheet_tokens.get(sheet) or set()
                score = 0
                for tok in list(toks)[:120]:
                    if tok and tok in ntext:
                        score += 1
                if score > best_score:
                    best_score = score
                    best_sheet = sheet
            return best_sheet
        scored: List[Tuple[str, int]] = []
        for sheet, toks in sheet_tokens.items():
            score = 0
            for tok in list(toks)[:180]:
                if tok and tok in ntext:
                    score += 1
            if score > 0:
                scored.append((sheet, score))
        if not scored:
            return None
        scored.sort(key=lambda x: x[1], reverse=True)
        best_sheet, best_score = scored[0]
        second_score = scored[1][1] if len(scored) > 1 else -1
        if best_score >= 6 and best_score >= second_score + 3:
            return best_sheet
        return None

    async def _extract_vlm_blocks_from_snapshot_images(
        self,
        file_name: str,
        ext: str,
        images: List[Dict[str, Any]],
        start_index: int,
        max_images: int = 10,
    ) -> List[Dict[str, Any]]:
        """
        对“整页快照图”执行 VLM 结构化解析。
        """
        if not settings.OPENAI_VISION_ENABLED or not settings.OPENAI_API_KEY:
            return []
        if not images:
            return []

        blocks: List[Dict[str, Any]] = []
        block_idx = start_index
        max_images = max(1, int(max_images or 1))
        for img in images[:max_images]:
            try:
                logger.info(
                    "VLM_CALL_START: file=%s, image=%s, sheet=%s, page=%s, storage_path=%s",
                    file_name,
                    img.get("name"),
                    img.get("sheet"),
                    img.get("page"),
                    img.get("storage_path"),
                )
                prompt = self._build_vlm_diagram_prompt(file_name=file_name, image_name=str(img.get("name") or "page.png"))
                payload = await self.chat_client.analyze_image_json(
                    image_bytes=img["bytes"],
                    mime_type=img.get("mime_type") or "image/png",
                    prompt=prompt,
                )
                if not payload:
                    logger.warning(
                        "VLM_CALL_EMPTY: file=%s, image=%s, sheet=%s, page=%s",
                        file_name,
                        img.get("name"),
                        img.get("sheet"),
                        img.get("page"),
                    )
                    continue
                logger.info(
                    "VLM_CALL_RESULT: file=%s, image=%s, payload_keys=%s",
                    file_name,
                    img.get("name"),
                    sorted(list(payload.keys())) if isinstance(payload, dict) else "non_dict",
                )
                logger.info(
                    "VLM_RAW_PAYLOAD_JSON: file=%s, image=%s, sheet=%s, page=%s, scope=sheet_snapshot, payload=%s",
                    file_name,
                    img.get("name"),
                    img.get("sheet"),
                    img.get("page"),
                    json.dumps(payload, ensure_ascii=False),
                )
                page_no = img.get("page")
                sheet_name = self._clean_text(img.get("sheet"))
                image_path = self._clean_text(img.get("storage_path"))
                section_name = str(img.get("name") or "")

                fallback_page_name = self._clean_text(img.get("sheet")) or self._clean_text(img.get("name")) or "unknown"
                fallback_page_id = (
                    f"{self._clean_text(img.get('sheet'))}_p{int(img.get('page') or 1)}"
                    if self._clean_text(img.get("sheet"))
                    else None
                )
                normalized = self._normalize_vlm_diagram_payload(
                    payload,
                    default_page_name=fallback_page_name,
                    default_page_id=fallback_page_id,
                )
                logger.info(
                    "VLM_OVERVIEW_FULL: file=%s, image=%s, sheet=%s, page=%s, scope=sheet_snapshot, diagram_overview=%s, process_description=%s, relationships=%s, evidence_texts=%s",
                    file_name,
                    img.get("name"),
                    sheet_name,
                    page_no,
                    self._clean_text(payload.get("diagram_overview") or payload.get("diagram_summary")),
                    self._clean_text(payload.get("process_description")),
                    json.dumps(self._clean_text_list(payload.get("relationships"), limit=80), ensure_ascii=False),
                    json.dumps(self._clean_text_list(payload.get("evidence_texts"), limit=80), ensure_ascii=False),
                )
                logger.info(
                    "VLM_NORMALIZED_FULL_SUMMARY: file=%s, image=%s, sheet=%s, page=%s, scope=sheet_snapshot, summary_text=%s",
                    file_name,
                    img.get("name"),
                    sheet_name,
                    page_no,
                    self._clean_text(normalized.get("summary_text")),
                )
                logger.info(
                    "VLM_NORMALIZED: file=%s, image=%s, sheet=%s, page=%s, summary_preview=%s, nodes=%s, edges=%s, confidence=%s",
                    file_name,
                    img.get("name"),
                    sheet_name,
                    page_no,
                    self._short_text(normalized.get("summary_text"), 120),
                    len(normalized.get("nodes", []) or []),
                    len(normalized.get("edges", []) or []),
                    normalized.get("confidence"),
                )
                page_id = self._clean_text(normalized.get("page_id")) or self._build_fallback_page_id(fallback_page_name, fallback_page_id)
                page_name = self._clean_text(normalized.get("page_name")) or fallback_page_name
                diagram_type = self._clean_text(normalized.get("diagram_type")) or "unknown"
                summary = self._clean_text(normalized.get("summary_text"))
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "diagram_page",
                        "text": (
                            f"[diagram_page] page_id={page_id} ; page_name={page_name} ; "
                            f"diagram_type={diagram_type}"
                        ),
                        "page": page_no,
                        "section": section_name or page_name,
                        "sheet": sheet_name or None,
                        "source_parser": "vlm_sheet_snapshot",
                        "file_type": ext,
                        "file_name": file_name,
                        "image_path": image_path or None,
                    }
                )
                block_idx += 1
                if summary:
                    blocks.append(
                        {
                            "block_index": block_idx,
                            "type": "diagram_summary",
                            "text": f"[diagram_summary] {summary}",
                            "page": page_no,
                            "section": section_name or page_name,
                            "sheet": sheet_name or None,
                            "source_parser": "vlm_sheet_snapshot",
                            "file_type": ext,
                            "file_name": file_name,
                            "image_path": image_path or None,
                        }
                    )
                    block_idx += 1

                for name in normalized.get("nodes", [])[:100]:
                    name = self._clean_text(name)
                    if not name:
                        continue
                    blocks.append(
                        {
                            "block_index": block_idx,
                            "type": "diagram_node",
                            "text": f"[diagram_node] {name}",
                            "page": page_no,
                            "section": section_name or page_name,
                            "sheet": sheet_name or None,
                            "source_parser": "vlm_sheet_snapshot",
                            "file_type": ext,
                            "file_name": file_name,
                            "image_path": image_path or None,
                        }
                    )
                    block_idx += 1

                for row in normalized.get("edges", [])[:140]:
                    src = self._clean_text(row.get("src"))
                    dst = self._clean_text(row.get("dst"))
                    if not src or not dst or src == dst:
                        continue
                    rel = self._clean_text(row.get("relation_type")) or "连接"
                    conf = self._coerce_confidence(row.get("confidence"), default=normalized.get("confidence", 0.70))
                    blocks.append(
                        {
                            "block_index": block_idx,
                            "type": "diagram_edge",
                            "text": f"{src} -> {dst} ; rel={rel} ; confidence={conf:.2f}",
                            "page": page_no,
                            "section": section_name or page_name,
                            "sheet": sheet_name or None,
                            "source_parser": "vlm_sheet_snapshot",
                            "file_type": ext,
                            "file_name": file_name,
                            "image_path": image_path or None,
                        }
                    )
                    block_idx += 1
            except Exception as e:
                logger.warning(
                    "VLM_CALL_ERROR: file=%s, image=%s, sheet=%s, page=%s, err=%s",
                    file_name,
                    img.get("name"),
                    img.get("sheet"),
                    img.get("page"),
                    e,
                )
                continue
        logger.info(
            "VLM_BLOCKS_TOTAL: file=%s, images=%s, generated_blocks=%s",
            file_name,
            min(len(images), 6),
            len(blocks),
        )
        return blocks

    async def _extract_vlm_blocks_from_xlsx_diagrams(
        self,
        file_name: str,
        graphs: Dict[str, Dict[str, Any]],
        start_index: int,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        logger.info(
            "XLSX重绘图VLM链路已禁用: file=%s, graphs=%s, reason=use_original_sheet_snapshots_only",
            file_name,
            len(graphs or {}),
        )
        return [], []

    def _should_enable_vlm_diagram_route(
        self,
        ext: str,
        quality: Dict[str, float],
        file_data: bytes,
        strategy: Optional[ProfileStrategy],
        blocks: Optional[List[Dict[str, Any]]] = None,
    ) -> bool:
        if not settings.OPENAI_VISION_ENABLED:
            return False
        if not settings.OPENAI_API_KEY:
            return False
        ext_norm = (ext or "").lower()
        if ext_norm not in ("docx", "xlsx", "xlsm", "xltx", "xltm"):
            return False
        profile_id = (strategy.profile_id if strategy else "general")
        profile_bias = 1 if profile_id in {"design", "ops"} else 0
        media_count = self._count_office_media_images(file_data, ext_norm)

        if ext_norm == "docx" and media_count <= 0:
            return False

        xlsx_visual_signals = False
        if ext_norm in ("xlsx", "xlsm", "xltx", "xltm"):
            candidate_sheets = self._collect_snapshot_candidate_sheets(blocks or [])
            xlsx_visual_signals = bool(
                candidate_sheets
                or float(quality.get("diagram_like_blocks", 0.0)) > 0
                or float(quality.get("image_like_blocks", 0.0)) > 0
            )
            if xlsx_visual_signals:
                return True

        if ext_norm == "docx" and media_count >= 1 and float(quality.get("total_chars", 0.0)) < 1500:
            return True

        weak_text = (
            float(quality.get("total_chars", 0.0)) < 1200
            or float(quality.get("avg_chars_per_block", 0.0)) < 40
        )
        structural_hits = (
            float(quality.get("table_like_blocks", 0.0))
            + float(quality.get("schedule_like_blocks", 0.0))
            + float(quality.get("diagram_like_blocks", 0.0))
        )
        score = 0
        if weak_text:
            score += 2
        if ext_norm == "docx" and media_count >= 2:
            score += 1
        if ext_norm in ("xlsx", "xlsm", "xltx", "xltm") and xlsx_visual_signals:
            score += 1
        if structural_hits < 2:
            score += 1
        score += profile_bias
        return score >= 2

    def _extract_with_tika(self, file_data: bytes, file_name: str) -> str:
        try:
            parsed = tika_parser.from_buffer(file_data)
            content = ""
            if parsed and "content" in parsed and parsed["content"]:
                content = self._clean_text(parsed["content"])
            if content:
                logger.info(f"Tika 解析成功（兜底）: {file_name}")
                return content
        except Exception as e:
            logger.warning(f"Tika 解析失败: {file_name}, err={e}")

        try:
            return self._clean_text(file_data.decode("utf-8"))
        except UnicodeDecodeError as e:
            raise ValueError(f"无法解析文件内容（Tika失败且非UTF-8）: {file_name}, err={e}") from e

    @staticmethod
    def _analyze_blocks_router_quality(blocks: List[Dict[str, Any]]) -> Dict[str, float]:
        """
        通用质量评估（用于“是否升级到 OCR/VLM 路由”判断）。
        """
        text_blocks = 0
        total_chars = 0
        table_like = 0
        schedule_like = 0
        diagram_like = 0
        image_like = 0
        for b in blocks or []:
            text = DocumentProcessorService._clean_text(b.get("text"))
            if text:
                text_blocks += 1
                total_chars += len(text)
            btype = str(b.get("type") or "")
            if "table" in btype:
                table_like += 1
            if "schedule" in btype:
                schedule_like += 1
            if "diagram" in btype:
                diagram_like += 1
            if "image" in btype:
                image_like += 1

        avg_chars = (total_chars / text_blocks) if text_blocks else 0.0
        return {
            "text_blocks": float(text_blocks),
            "total_chars": float(total_chars),
            "avg_chars_per_block": float(avg_chars),
            "table_like_blocks": float(table_like),
            "schedule_like_blocks": float(schedule_like),
            "diagram_like_blocks": float(diagram_like),
            "image_like_blocks": float(image_like),
        }

    @staticmethod
    def _count_office_media_images(file_data: bytes, ext: str) -> int:
        """
        统计 office 文档内嵌图片数量（docx/xlsx 通用，基于 zip media 路径）。
        """
        try:
            zf = zipfile.ZipFile(BytesIO(file_data))
        except Exception:
            return 0
        ext_norm = ext.lower()
        if ext_norm == "docx":
            prefix = "word/media/"
        elif ext_norm in ("xlsx", "xlsm", "xltx", "xltm"):
            prefix = "xl/media/"
        else:
            return 0
        return sum(1 for n in zf.namelist() if n.startswith(prefix))

    def _extract_office_image_ocr_blocks(
        self,
        file_data: bytes,
        file_name: str,
        ext: str,
        start_index: int,
        max_images: int = 8,
    ) -> List[Dict[str, Any]]:
        """
        Office（docx/xlsx）图片 OCR 兜底。
        仅在质量门控触发时使用，避免常态高成本。
        """
        if pytesseract is None:
            return []
        try:
            zf = zipfile.ZipFile(BytesIO(file_data))
        except Exception:
            return []

        ext_norm = ext.lower()
        if ext_norm == "docx":
            prefix = "word/media/"
        elif ext_norm in ("xlsx", "xlsm", "xltx", "xltm"):
            prefix = "xl/media/"
        else:
            return []

        image_names = sorted([n for n in zf.namelist() if n.startswith(prefix)])[:max_images]
        blocks: List[Dict[str, Any]] = []
        block_idx = start_index
        for i, name in enumerate(image_names, start=1):
            try:
                img_bytes = zf.read(name)
                if not img_bytes:
                    continue
                img = Image.open(BytesIO(img_bytes))
                ocr_text = self._clean_text(pytesseract.image_to_string(img, lang="jpn+eng"))
                if len(ocr_text) < 10:
                    continue
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "office_image_ocr",
                        "text": ocr_text,
                        "page": None,
                        "section": f"image_{i}",
                        "sheet": None,
                        "source_parser": "office_ocr",
                        "file_type": ext_norm,
                        "file_name": file_name,
                    }
                )
                block_idx += 1
            except Exception as e:
                logger.debug(f"Office 图片 OCR 失败: file={file_name}, image={name}, err={e}")
                continue
        return blocks

    def _should_escalate_to_vision_route(
        self,
        ext: str,
        quality: Dict[str, float],
        file_data: bytes,
    ) -> bool:
        """
        通用“是否升级到 OCR/VLM 路由”判断。
        规则：
        - 文本质量弱（总字数/平均字数偏低）
        - 或者存在较多内嵌图片且结构化命中弱
        """
        ext_norm = ext.lower()
        if ext_norm not in ("docx", "xlsx", "xlsm", "xltx", "xltm"):
            return False
        weak_text = (
            quality.get("total_chars", 0.0) < 1200
            or quality.get("avg_chars_per_block", 0.0) < 40
        )
        has_schedule = quality.get("schedule_like_blocks", 0.0) > 0
        media_count = self._count_office_media_images(file_data, ext_norm)
        visual_heavy = media_count >= 2
        return weak_text or (visual_heavy and not has_schedule)

    @staticmethod
    def _analyze_page_text_quality(text: str) -> Dict[str, float]:
        normalized = DocumentProcessorService._clean_text(text)
        chars = len(normalized)
        words = re.findall(r"[\u3040-\u30ff\u3400-\u9fffA-Za-z0-9]{2,}", normalized)
        valid_tokens = [
            w for w in words if re.search(r"[\u3040-\u30ff\u3400-\u9fffA-Za-z]", w)
        ]
        valid_ratio = (len(valid_tokens) / len(words)) if words else 0.0
        return {"chars": float(chars), "words": float(len(words)), "valid_ratio": float(valid_ratio)}

    @staticmethod
    def _should_ocr_page(
        quality: Dict[str, float],
        no_text_streak: int,
        strategy: ProfileStrategy,
        doc_force_ocr: bool,
    ) -> bool:
        low_chars = quality["chars"] < strategy.ocr_min_chars_per_page
        low_words = quality["words"] < strategy.ocr_min_words_per_page
        low_valid = quality["valid_ratio"] < strategy.ocr_min_valid_ratio
        streak_triggered = no_text_streak >= strategy.ocr_no_text_streak_trigger
        return doc_force_ocr or low_chars or low_words or low_valid or streak_triggered

    def _parse_pdf_blocks(
        self,
        file_data: bytes,
        file_name: str,
        profile_strategy: Optional[ProfileStrategy] = None,
    ) -> List[Dict[str, Any]]:
        """
        PDF 解析：
        1) 优先提取文本层（按页）
        2) 对文本层缺失页，尝试 OCR（若环境支持）
        3) 额外抽取 PDF 表格
        4) 对文本层极弱/缺失场景，使用 Tika 兜底
        """
        blocks: List[Dict[str, Any]] = []
        block_idx = 0

        reader = PdfReader(BytesIO(file_data))
        strategy = profile_strategy or profile_service.get_strategy("general")
        page_items: List[Dict[str, Any]] = []
        weak_pages = 0
        no_text_streak = 0

        for page_idx, page in enumerate(reader.pages, start=1):
            raw_text = page.extract_text() or ""
            text = self._clean_text(raw_text)
            quality = self._analyze_page_text_quality(text)
            if quality["chars"] > 0:
                no_text_streak = 0
            else:
                no_text_streak += 1

            weak_base = (
                quality["chars"] < strategy.ocr_min_chars_per_page
                or quality["words"] < strategy.ocr_min_words_per_page
                or quality["valid_ratio"] < strategy.ocr_min_valid_ratio
                or no_text_streak >= strategy.ocr_no_text_streak_trigger
            )
            if weak_base:
                weak_pages += 1

            page_items.append(
                {
                    "page_idx": page_idx,
                    "text": text,
                    "quality": quality,
                    "no_text_streak": no_text_streak,
                }
            )

        total_pages = max(1, len(page_items))
        doc_force_ocr = strategy.enable_adaptive_ocr and (weak_pages / total_pages >= strategy.ocr_doc_trigger_ratio)
        ocr_budget = strategy.ocr_max_pages if strategy.enable_adaptive_ocr else 0
        ocr_used = 0
        pages_with_text = 0

        for item in page_items:
            page_idx = int(item["page_idx"])
            text = str(item["text"] or "")
            quality = item["quality"]
            no_text_streak = int(item["no_text_streak"])
            should_ocr = strategy.enable_adaptive_ocr and self._should_ocr_page(
                quality=quality,
                no_text_streak=no_text_streak,
                strategy=strategy,
                doc_force_ocr=doc_force_ocr,
            )

            ocr_text = ""
            if should_ocr and ocr_used < ocr_budget:
                ocr_text = self._ocr_pdf_page(file_data=file_data, page_number=page_idx)
                if ocr_text:
                    ocr_used += 1
                    blocks.append(
                        {
                            "block_index": block_idx,
                            "type": "pdf_page_ocr",
                            "text": ocr_text,
                            "page": page_idx,
                            "section": None,
                            "sheet": None,
                            "source_parser": "pdf_ocr",
                            "file_type": "pdf",
                            "file_name": file_name,
                        }
                    )
                    block_idx += 1
                    continue

            if text:
                pages_with_text += 1
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "pdf_page_text",
                        "text": text,
                        "page": page_idx,
                        "section": None,
                        "sheet": None,
                        "source_parser": "pdf_text_layer",
                        "file_type": "pdf",
                        "file_name": file_name,
                    }
                )
                block_idx += 1
            else:
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "pdf_image_page",
                        "text": f"[PAGE_{page_idx}] 该页未检测到可用文本层，建议OCR处理。",
                        "page": page_idx,
                        "section": None,
                        "sheet": None,
                        "source_parser": "pdf_text_layer",
                        "file_type": "pdf",
                        "file_name": file_name,
                    }
                )
                block_idx += 1

        table_blocks = self._extract_pdf_table_blocks(file_data=file_data, file_name=file_name, start_index=block_idx)
        if table_blocks:
            blocks.extend(table_blocks)
            block_idx = blocks[-1]["block_index"] + 1

        text_ratio = pages_with_text / max(1, len(reader.pages))
        if not blocks or text_ratio < 0.3:
            fallback_text = self._extract_with_tika(file_data, file_name)
            if fallback_text:
                blocks.append(
                    {
                        "block_index": block_idx,
                        "type": "pdf_tika_fallback",
                        "text": fallback_text,
                        "page": None,
                        "section": None,
                        "sheet": None,
                        "source_parser": "tika",
                        "file_type": "pdf",
                        "file_name": file_name,
                    }
                )

        logger.info(
            "PDF解析路由: file=%s, total_pages=%s, weak_pages=%s, doc_force_ocr=%s, ocr_used=%s/%s, text_ratio=%.2f",
            file_name,
            len(reader.pages),
            weak_pages,
            doc_force_ocr,
            ocr_used,
            ocr_budget,
            text_ratio,
        )
        return blocks

    def _ocr_pdf_page(self, file_data: bytes, page_number: int) -> str:
        if fitz is None or pytesseract is None:
            return ""

        try:
            doc = fitz.open(stream=file_data, filetype="pdf")
            page = doc.load_page(page_number - 1)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            image = Image.open(BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(image, lang="jpn+eng")
            return self._clean_text(text)
        except Exception as e:
            logger.debug(f"PDF OCR 失败: page={page_number}, err={e}")
            return ""

    def _extract_pdf_table_blocks(
        self, file_data: bytes, file_name: str, start_index: int
    ) -> List[Dict[str, Any]]:
        if pdfplumber is None:
            return []

        blocks: List[Dict[str, Any]] = []
        block_idx = start_index

        try:
            with pdfplumber.open(BytesIO(file_data)) as pdf:
                for page_idx, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables() or []
                    for table_no, table in enumerate(tables, start=1):
                        if not table:
                            continue
                        header: Optional[List[str]] = None
                        for row_idx, row in enumerate(table, start=1):
                            if row is None:
                                continue
                            values = [self._clean_text(v) for v in row]
                            if not any(values):
                                continue
                            if header is None:
                                header = [v or f"col_{i+1}" for i, v in enumerate(values)]
                                text = " | ".join(header)
                                block_type = "pdf_table_header"
                                row_json = {f"col_{i+1}": v for i, v in enumerate(values)}
                            else:
                                pairs = []
                                row_json = {}
                                for i, value in enumerate(values):
                                    if not value:
                                        continue
                                    col_name = header[i] if i < len(header) else f"col_{i+1}"
                                    pairs.append(f"{col_name}: {value}")
                                    row_json[col_name] = value
                                text = " ; ".join(pairs) if pairs else " | ".join(values)
                                block_type = "pdf_table_row"

                            blocks.append(
                                {
                                    "block_index": block_idx,
                                    "type": block_type,
                                    "text": text,
                                    "page": page_idx,
                                    "section": None,
                                    "sheet": f"table_{table_no}",
                                    "table_name": f"table_{table_no}",
                                    "row_no": row_idx,
                                    "row_json": row_json,
                                    "source_parser": "pdf_table",
                                    "file_type": "pdf",
                                    "file_name": file_name,
                                }
                            )
                            block_idx += 1
        except Exception as e:
            logger.warning(f"PDF 表格抽取失败: file={file_name}, err={e}")

        return blocks

    def _parse_generic_blocks(self, file_data: bytes, file_name: str, file_type: str) -> List[Dict[str, Any]]:
        text = self._extract_with_tika(file_data, file_name)
        if not text:
            return []
        return [
            {
                "block_index": 0,
                "type": "text",
                "text": text,
                "page": None,
                "section": None,
                "sheet": None,
                "source_parser": "tika",
                "file_type": file_type,
                "file_name": file_name,
            }
        ]

    def parse_document_blocks(
        self,
        file_data: bytes,
        file_name: str,
        profile_strategy: Optional[ProfileStrategy] = None,
    ) -> List[Dict[str, Any]]:
        """
        按文件类型路由解析，产出结构化 blocks。
        """
        ext = Path(file_name).suffix.lower().lstrip(".")
        if ext in ("docx",):
            blocks = self._parse_docx_blocks(file_data, file_name)
        elif ext in ("xlsx", "xlsm", "xltx", "xltm"):
            blocks = self._parse_xlsx_blocks(file_data, file_name)
        elif ext == "pdf":
            blocks = self._parse_pdf_blocks(file_data, file_name, profile_strategy=profile_strategy)
        else:
            blocks = self._parse_generic_blocks(file_data, file_name, ext or "unknown")

        quality = self._analyze_blocks_router_quality(blocks)
        if self._should_escalate_to_vision_route(ext, quality, file_data):
            start_idx = (max((int(b.get("block_index") or 0) for b in blocks), default=-1) + 1)
            ocr_blocks = self._extract_office_image_ocr_blocks(
                file_data=file_data,
                file_name=file_name,
                ext=ext,
                start_index=start_idx,
                max_images=8,
            )
            if ocr_blocks:
                blocks.extend(ocr_blocks)
                logger.info(
                    "文档解析路由升级: file=%s, ext=%s, reason=quality_gate, office_ocr_blocks=%s, quality=%s",
                    file_name,
                    ext,
                    len(ocr_blocks),
                    quality,
                )
            elif quality.get("total_chars", 0.0) < 800:
                fallback_text = self._extract_with_tika(file_data, file_name)
                if fallback_text:
                    blocks.append(
                        {
                            "block_index": start_idx,
                            "type": "quality_tika_fallback",
                            "text": fallback_text,
                            "page": None,
                            "section": None,
                            "sheet": None,
                            "source_parser": "tika",
                            "file_type": ext or "unknown",
                            "file_name": file_name,
                        }
                    )
                    logger.info(
                        "文档解析路由升级: file=%s, ext=%s, reason=quality_gate, fallback=tika, quality=%s",
                        file_name,
                        ext,
                        quality,
                    )

        if ext in ("md", "markdown"):
            for block in blocks:
                block["text"] = self._clean_markdown(block["text"])

        valid_blocks = [b for b in blocks if self._clean_text(b.get("text"))]
        logger.info(f"结构化解析完成: {file_name}, blocks={len(valid_blocks)}")
        return valid_blocks

    def _clean_markdown(self, text: str) -> str:
        text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"```[^`]*```", "", text, flags=re.DOTALL)
        text = re.sub(r"`[^`]+`", "", text)
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
        text = re.sub(r"\*\*([^\*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^\*]+)\*", r"\1", text)
        return self._clean_text(text)

    def _split_long_text(
        self,
        text: str,
        base_meta: Dict[str, Any],
        start_chunk_id: int,
    ) -> List[Dict[str, Any]]:
        chunks: List[Dict[str, Any]] = []
        start = 0
        chunk_id = start_chunk_id
        max_iterations = len(text) // max(1, self.chunk_size - self.chunk_overlap) + 10
        iteration = 0

        while start < len(text) and iteration < max_iterations:
            iteration += 1
            end = min(start + self.chunk_size, len(text))
            chunk_text = self._clean_text(text[start:end])
            if chunk_text:
                chunks.append(
                    {
                        "chunk_id": chunk_id,
                        "text": chunk_text,
                        "metadata": dict(base_meta),
                        "sources": [
                            {
                                "block_index": base_meta.get("block_index"),
                                "source_type": base_meta.get("chunk_type"),
                                "source_parser": base_meta.get("source_parser"),
                                "sheet": base_meta.get("sheet"),
                                "page": base_meta.get("page"),
                                "section": base_meta.get("section"),
                                "image_path": base_meta.get("image_path"),
                                "text_preview": chunk_text[:500],
                            }
                        ],
                    }
                )
                chunk_id += 1

            next_start = end - self.chunk_overlap
            if next_start <= start:
                next_start = start + max(1, self.chunk_size - self.chunk_overlap)
            start = next_start

        return chunks

    def split_blocks_into_chunks(self, blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        结构优先切块：
        1) 优先按 block 聚合
        2) 超长文本再滑窗兜底
        """
        if not blocks:
            return []

        chunks: List[Dict[str, Any]] = []
        current_parts: List[str] = []
        current_meta: Dict[str, Any] = {}
        current_sources: List[Dict[str, Any]] = []
        chunk_id = 0

        def flush_current() -> None:
            nonlocal chunk_id, current_parts, current_meta, current_sources
            if not current_parts:
                return
            text = self._clean_text("\n\n".join(current_parts))
            if text:
                chunks.append(
                    {
                        "chunk_id": chunk_id,
                        "text": text,
                        "metadata": dict(current_meta),
                        "sources": list(current_sources),
                    }
                )
                chunk_id += 1
            current_parts = []
            current_meta = {}
            current_sources = []

        for block in blocks:
            text = self._clean_text(block.get("text"))
            if not text:
                continue

            meta = {
                "block_index": block.get("block_index"),
                "chunk_type": block.get("type"),
                "page": block.get("page"),
                "section": block.get("section"),
                "sheet": block.get("sheet"),
                "source_parser": block.get("source_parser"),
                "file_type": block.get("file_type"),
                "image_path": block.get("image_path"),
            }
            source_ref = {
                "block_index": block.get("block_index"),
                "source_type": block.get("type"),
                "source_parser": block.get("source_parser"),
                "sheet": block.get("sheet"),
                "page": block.get("page"),
                "section": block.get("section"),
                "image_path": block.get("image_path"),
                "text_preview": text[:500],
            }

            current_sheet = str(current_meta.get("sheet") or "").strip()
            incoming_sheet = str(meta.get("sheet") or "").strip()
            current_file_type = str(current_meta.get("file_type") or "").strip().lower()
            incoming_file_type = str(meta.get("file_type") or "").strip().lower()
            if (
                current_parts
                and current_file_type == "xlsx"
                and incoming_file_type == "xlsx"
                and current_sheet
                and incoming_sheet
                and current_sheet != incoming_sheet
            ):
                flush_current()

            if len(text) > self.chunk_size * 2:
                flush_current()
                for split_chunk in self._split_long_text(text, meta, chunk_id):
                    chunks.append(split_chunk)
                if chunks:
                    chunk_id = chunks[-1]["chunk_id"] + 1
                continue

            pending_text = "\n\n".join(current_parts + [text]) if current_parts else text
            if len(pending_text) > self.chunk_size:
                flush_current()
                current_parts = [text]
                current_meta = dict(meta)
                current_sources = [source_ref]
            else:
                current_parts.append(text)
                current_sources.append(source_ref)
                if not current_meta:
                    current_meta = dict(meta)
                elif current_meta.get("chunk_type") != meta.get("chunk_type"):
                    current_meta["chunk_type"] = "mixed"
                if current_meta.get("page") is None and meta.get("page") is not None:
                    current_meta["page"] = meta.get("page")
                if not current_meta.get("section") and meta.get("section"):
                    current_meta["section"] = meta.get("section")
                if not current_meta.get("sheet") and meta.get("sheet"):
                    current_meta["sheet"] = meta.get("sheet")
                if not current_meta.get("image_path") and meta.get("image_path"):
                    current_meta["image_path"] = meta.get("image_path")

        flush_current()
        logger.info(f"结构化切块完成: chunks={len(chunks)}")
        return chunks

    def _collect_table_rows(self, blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        rows: List[Dict[str, Any]] = []
        for block in blocks:
            block_type = str(block.get("type") or "")
            if "table" not in block_type and block_type != "schedule_row":
                continue

            row_json = block.get("row_json")
            if not isinstance(row_json, dict) or not row_json:
                row_json = {}
                text = self._clean_text(block.get("text"))
                if text:
                    for seg in [s.strip() for s in text.split(";") if s.strip()]:
                        if ":" in seg:
                            k, v = seg.split(":", 1)
                            key = self._clean_text(k)
                            val = self._clean_text(v)
                            if key and val:
                                row_json[key] = val
                if not row_json:
                    continue

            row_no_raw = block.get("row_no") or block.get("row")
            try:
                row_no = int(row_no_raw) if row_no_raw is not None else None
            except Exception:
                row_no = None

            rows.append(
                {
                    "sheet": block.get("sheet"),
                    "table_name": block.get("table_name") or block.get("sheet"),
                    "row_no": row_no,
                    "row_json": row_json,
                    "raw_text": self._clean_text(block.get("text")),
                    "source_parser": block.get("source_parser") or "table",
                }
            )
        return rows

    async def process_document(
        self,
        file_md5: str,
        file_name: str,
        storage_path: str,
        user_id: int,
        org_tag: Optional[str] = None,
        is_public: bool = False,
        kb_profile: Optional[str] = None,
    ) -> bool:
        """
        处理文档：下载、解析、向量化、索引
        """
        if not db_client.SessionLocal:
            logger.error("数据库会话工厂未初始化")
            return False

        try:
            async with db_client.SessionLocal() as db:
                logger.info(f"开始处理文档: file_md5={file_md5}, file_name={file_name}")
                selected_profile = kb_profile or await profile_service.get_selected_profile(db) or "general"
                profile_strategy = profile_service.get_strategy(selected_profile)
                self.chunk_size = int(getattr(settings, "CHUNK_SIZE", self.chunk_size))
                self.chunk_overlap = int(getattr(settings, "CHUNK_OVERLAP", self.chunk_overlap))
                if self.chunk_overlap >= self.chunk_size:
                    self.chunk_overlap = max(0, self.chunk_size // 4)

                file_data = minio_client.download_file(
                    bucket_name=settings.MINIO_DEFAULT_BUCKET,
                    object_name=storage_path
                )
                if not file_data:
                    logger.error(f"文件下载失败: {storage_path}")
                    return False

                blocks = self.parse_document_blocks(
                    file_data,
                    file_name,
                    profile_strategy=profile_strategy,
                )
                if not blocks:
                    logger.warning(f"结构化解析结果为空: {file_name}")
                    return False

                ext = Path(file_name).suffix.lower().lstrip(".")
                quality = self._analyze_blocks_router_quality(blocks)
                max_vlm_images = max(1, int(getattr(settings, "OPENAI_VISION_MAX_IMAGES_PER_FILE", 10) or 10))
                max_snapshot_pages = max(1, int(getattr(settings, "XLSX_SNAPSHOT_MAX_PAGES_PER_SHEET", 3) or 3))
                should_enable_vlm = self._should_enable_vlm_diagram_route(
                    ext=ext,
                    quality=quality,
                    file_data=file_data,
                    strategy=profile_strategy,
                    blocks=blocks,
                )
                logger.info(
                    "VLM_ROUTE_DECISION: file=%s, ext=%s, profile=%s, should_enable=%s, quality=%s, limits={max_images=%s,max_pages_per_sheet=%s}, vision_enabled=%s, has_api_key=%s",
                    file_name,
                    ext,
                    selected_profile,
                    should_enable_vlm,
                    {
                        "text_blocks": int(quality.get("text_blocks", 0)),
                        "total_chars": int(quality.get("total_chars", 0)),
                        "avg_chars_per_block": round(float(quality.get("avg_chars_per_block", 0.0)), 2),
                        "table_like_blocks": int(quality.get("table_like_blocks", 0)),
                        "schedule_like_blocks": int(quality.get("schedule_like_blocks", 0)),
                        "diagram_like_blocks": int(quality.get("diagram_like_blocks", 0)),
                        "image_like_blocks": int(quality.get("image_like_blocks", 0)),
                    },
                    max_vlm_images,
                    max_snapshot_pages,
                    settings.OPENAI_VISION_ENABLED,
                    bool(settings.OPENAI_API_KEY),
                )
                if should_enable_vlm:
                    if ext in ("xlsx", "xlsm", "xltx", "xltm"):
                        logger.info(
                            "XLSX VLM deferred: file=%s, reason=run_after_sheet_snapshots",
                            file_name,
                        )
                    else:
                        next_idx = (max([int(b.get("block_index") or 0) for b in blocks]) + 1) if blocks else 0
                        vlm_blocks, office_vlm_images = await self._extract_office_vlm_diagram_blocks(
                            file_data=file_data,
                            file_name=file_name,
                            file_md5=file_md5,
                            user_id=user_id,
                            ext=ext,
                            start_index=next_idx,
                            max_images=max_vlm_images,
                        )
                        await db.execute(
                            ImageBlock.__table__.delete().where(
                                ImageBlock.file_md5 == file_md5,
                                ImageBlock.source_parser == "vlm_diagram",
                            )
                        )
                        if office_vlm_images:
                            for item in office_vlm_images:
                                db.add(
                                    ImageBlock(
                                        file_md5=file_md5,
                                        sheet=item.get("sheet"),
                                        page=item.get("page"),
                                        anchor_row=item.get("anchor_row"),
                                        anchor_col=item.get("anchor_col"),
                                        storage_path=item.get("storage_path"),
                                        image_md5=item.get("image_md5"),
                                        image_width=item.get("image_width"),
                                        image_height=item.get("image_height"),
                                        content_type=item.get("content_type"),
                                        description=item.get("description"),
                                        source_parser=item.get("source_parser") or "vlm_diagram",
                                    )
                                )
                            await db.flush()
                            logger.info(
                                "Office VLM image blocks persisted: file=%s, file_md5=%s, count=%s",
                                file_name,
                                file_md5,
                                len(office_vlm_images),
                            )
                        if vlm_blocks:
                            blocks.extend(vlm_blocks)
                            logger.info(
                                "VLM图形补强已启用: file=%s, ext=%s, profile=%s, added_blocks=%s",
                                file_name,
                                ext,
                                selected_profile,
                                len(vlm_blocks),
                            )

                xlsx_sheet_snapshots_for_vlm: List[Dict[str, Any]] = []


                table_rows_raw = self._collect_table_rows(blocks)
                try:
                    await db.execute(TableRow.__table__.delete().where(TableRow.file_md5 == file_md5))
                    if table_rows_raw:
                        for item in table_rows_raw:
                            db.add(
                                TableRow(
                                    file_md5=file_md5,
                                    sheet=item.get("sheet"),
                                    table_name=item.get("table_name"),
                                    row_no=item.get("row_no"),
                                    row_json=json.dumps(item.get("row_json", {}), ensure_ascii=False),
                                    raw_text=item.get("raw_text"),
                                    source_parser=item.get("source_parser") or "table",
                                )
                            )
                        await db.flush()
                    logger.info("表格行结构化入库完成: file=%s, rows=%s", file_name, len(table_rows_raw))
                except Exception as table_err:
                    logger.warning(f"表格行结构化入库失败（降级继续）: file={file_name}, err={table_err}")

                image_blocks_raw: List[Dict[str, Any]] = []
                if ext in ("xlsx", "xlsm", "xltx", "xltm"):
                    try:
                        await db.execute(
                            ImageBlock.__table__.delete().where(ImageBlock.file_md5 == file_md5)
                        )
                        image_blocks_raw = self._extract_xlsx_images(
                            file_data=file_data,
                            file_name=file_name,
                            file_md5=file_md5,
                            user_id=user_id,
                        )
                        logger.info(
                            "XLSX embedded images extracted: file=%s, file_md5=%s, count=%s",
                            file_name,
                            file_md5,
                            len(image_blocks_raw),
                        )
                        candidate_sheets = sorted(self._collect_snapshot_candidate_sheets(blocks))
                        logger.info(
                            "XLSX snapshot plan: file=%s, file_md5=%s, candidate_sheets=%s",
                            file_name,
                            file_md5,
                            candidate_sheets,
                        )
                        sheet_snapshots = self._render_xlsx_sheets_via_libreoffice(
                            file_data=file_data,
                            file_name=file_name,
                            target_sheets=candidate_sheets,
                            blocks=blocks,
                            max_pages_per_sheet=max_snapshot_pages,
                        )
                        if sheet_snapshots:
                            logger.info(
                                "XLSX snapshots generated: file=%s, file_md5=%s, count=%s, sheets=%s",
                                file_name,
                                file_md5,
                                len(sheet_snapshots),
                                sorted(
                                    {
                                        str(s.get("sheet") or "").strip()
                                        for s in sheet_snapshots
                                        if str(s.get("sheet") or "").strip()
                                    }
                                ),
                            )
                            snap_idx = 800
                            for snap in sheet_snapshots:
                                snap_idx += 1
                                raw_bytes = snap.get("bytes") or b""
                                if not raw_bytes:
                                    continue
                                match_mode = str(snap.get("match_mode") or "").strip().lower()
                                page_no = int(snap.get("page") or 0) or 1
                                mapped_sheet = str(snap.get("sheet") or "").strip()
                                object_sheet_name = f"snapshot_{mapped_sheet}" if mapped_sheet else "snapshot_unbound"
                                object_path = minio_client.build_document_image_path(
                                    user_id=user_id,
                                    file_md5=file_md5,
                                    sheet_name=object_sheet_name,
                                    image_index=snap_idx,
                                    ext="png",
                                )
                                uploaded = minio_client.upload_bytes(
                                    bucket_name=settings.MINIO_DEFAULT_BUCKET,
                                    object_name=object_path,
                                    data=raw_bytes,
                                    content_type="image/png",
                                )
                                if not uploaded:
                                    continue
                                img_w = 0
                                img_h = 0
                                try:
                                    with Image.open(BytesIO(raw_bytes)) as im:
                                        img_w, img_h = im.size
                                except Exception:
                                    pass
                                image_blocks_raw.append(
                                    {
                                        "sheet": mapped_sheet or None,
                                        "page": page_no,
                                        "anchor_row": None,
                                        "anchor_col": None,
                                        "storage_path": object_path,
                                        "image_md5": hashlib.md5(raw_bytes).hexdigest(),
                                        "image_width": img_w or None,
                                        "image_height": img_h or None,
                                        "content_type": "image/png",
                                        "description": (
                                            f"original workbook snapshot page={page_no}, "
                                            f"sheet={(mapped_sheet or 'unbound')}"
                                        ),
                                        "source_parser": "xlsx_sheet_snapshot",
                                        "match_mode": snap.get("match_mode"),
                                        "match_confidence": snap.get("match_confidence"),
                                        "match_reason": snap.get("match_reason"),
                                    }
                                )
                                xlsx_sheet_snapshots_for_vlm.append(
                                    {
                                        "name": snap.get("name") or f"{(mapped_sheet or 'unbound')}_page_{page_no}.png",
                                        "bytes": raw_bytes,
                                        "mime_type": "image/png",
                                        "page": page_no,
                                        "sheet": mapped_sheet or None,
                                        "storage_path": object_path,
                                    }
                                )
                            logger.info(
                                "XLSX snapshots appended to image_blocks_raw: file=%s, file_md5=%s, total_image_blocks=%s",
                                file_name,
                                file_md5,
                                len(image_blocks_raw),
                            )
                            logger.info(
                                "XLSX 原始页面快照保存: file=%s, sheets=%s, images=%s",
                                file_name,
                                len(candidate_sheets),
                                len(sheet_snapshots),
                            )
                        else:
                            logger.info(
                                "XLSX snapshots empty: file=%s, file_md5=%s, candidate_sheets=%s",
                                file_name,
                                file_md5,
                                candidate_sheets,
                            )
                        if image_blocks_raw:
                            for item in image_blocks_raw:
                                db.add(
                                    ImageBlock(
                                        file_md5=file_md5,
                                        sheet=item.get("sheet"),
                                        page=item.get("page"),
                                        anchor_row=item.get("anchor_row"),
                                        anchor_col=item.get("anchor_col"),
                                        storage_path=item.get("storage_path"),
                                        image_md5=item.get("image_md5"),
                                        image_width=item.get("image_width"),
                                        image_height=item.get("image_height"),
                                        content_type=item.get("content_type"),
                                        description=item.get("description"),
                                        source_parser=item.get("source_parser") or "xlsx_image",
                                        match_mode=item.get("match_mode"),
                                        match_confidence=(
                                            int(round(float(item.get("match_confidence") or 0.0) * 100))
                                            if item.get("match_confidence") is not None
                                            else None
                                        ),
                                        match_reason=item.get("match_reason"),
                                    )
                                )
                            await db.flush()
                            logger.info(
                                "XLSX image blocks persisted: file=%s, file_md5=%s, count=%s",
                                file_name,
                                file_md5,
                                len(image_blocks_raw),
                            )

                            next_idx = (max([int(b.get("block_index") or 0) for b in blocks]) + 1) if blocks else 0
                            for img_item in image_blocks_raw:
                                blocks.append(
                                    {
                                        "block_index": next_idx,
                                        "type": "xlsx_image",
                                        "text": (
                                            f"[图片] sheet={img_item.get('sheet') or '-'} "
                                            f"page={img_item.get('page') or '-'} "
                                            f"anchor=R{img_item.get('anchor_row') or '-'}C{img_item.get('anchor_col') or '-'} ; "
                                            f"description={img_item.get('description') or '-'}"
                                        ),
                                        "page": None,
                                        "section": None,
                                        "sheet": img_item.get("sheet"),
                                        "source_parser": img_item.get("source_parser") or "xlsx_image",
                                        "file_type": "xlsx",
                                        "file_name": file_name,
                                        "image_path": img_item.get("storage_path"),
                                    }
                                )
                                next_idx += 1
                        logger.info("XLSX 图片抽取完成: file=%s, images=%s", file_name, len(image_blocks_raw))
                    except Exception as img_e:
                        logger.warning(f"XLSX 图片抽取失败（降级继续）: file={file_name}, err={img_e}")

                if (
                    should_enable_vlm
                    and ext in ("xlsx", "xlsm", "xltx", "xltm")
                    and xlsx_sheet_snapshots_for_vlm
                ):
                    next_idx = (max([int(b.get("block_index") or 0) for b in blocks]) + 1) if blocks else 0
                    snapshot_vlm_blocks = await self._extract_vlm_blocks_from_snapshot_images(
                        file_name=file_name,
                        ext=ext,
                        images=xlsx_sheet_snapshots_for_vlm,
                        start_index=next_idx,
                        max_images=max_vlm_images,
                    )
                    if snapshot_vlm_blocks:
                        blocks.extend(snapshot_vlm_blocks)
                        logger.info(
                            "XLSX sheet快照VLM补强: file=%s, snapshots=%s, added_blocks=%s",
                            file_name,
                            len(xlsx_sheet_snapshots_for_vlm),
                            len(snapshot_vlm_blocks),
                        )
                    else:
                        logger.warning(
                            "XLSX_SNAPSHOT_VLM_EMPTY: file=%s, snapshots=%s, reason=no_blocks_generated",
                            file_name,
                            len(xlsx_sheet_snapshots_for_vlm),
                        )

                try:
                    if relation_search_service.should_build_relation_index(
                        file_name=file_name, blocks=blocks
                    ):
                        relation_stats = await relation_search_service.build_relation_index(
                            db=db,
                            file_md5=file_md5,
                            file_name=file_name,
                            blocks=blocks,
                        )
                        logger.info(
                            "关系索引构建: file_md5=%s, nodes=%s, edges=%s",
                            file_md5,
                            relation_stats.get("nodes", 0),
                            relation_stats.get("edges", 0),
                        )
                    else:
                        logger.info("文档未命中关系索引条件，跳过关系抽取: %s", file_name)
                except Exception as rel_err:
                    logger.warning(f"关系索引构建失败（已降级继续）: file={file_name}, err={rel_err}")

                chunks = self.split_blocks_into_chunks(blocks)
                if not chunks:
                    logger.warning(f"文本分块为空: {file_name}")
                    return False

                try:
                    exp_stats = await experience_service.rebuild_for_file(
                        db=db,
                        file_md5=file_md5,
                        chunks=chunks,
                    )
                    logger.info("经历索引构建: file_md5=%s, items=%s", file_md5, exp_stats.get("items", 0))
                except Exception as exp_err:
                    logger.warning(f"经历索引构建失败（已降级继续）: file={file_name}, err={exp_err}")

                await search_service.ensure_index_exists()

                texts = [chunk["text"] for chunk in chunks]
                vectors = await embedding_service.embed_batch(texts)
                successful_vectors = sum(1 for v in vectors if v is not None)
                logger.info(f"向量化完成: {successful_vectors}/{len(chunks)}")
                if successful_vectors == 0:
                    logger.error("所有文本块向量化失败")
                    return False

                file_result = await db.execute(
                    select(FileUpload).where(
                        FileUpload.file_md5 == file_md5,
                        FileUpload.user_id == user_id
                    )
                )
                file_record = file_result.scalar_one_or_none()
                if not file_record:
                    logger.error(f"文件记录不存在: file_md5={file_md5}, user_id={user_id}")
                    return False
                if file_record.kb_profile:
                    selected_profile = file_record.kb_profile

                user_result = await db.execute(select(User).where(User.id == user_id))
                user = user_result.scalar_one_or_none()
                if not user:
                    logger.error(f"用户不存在: user_id={user_id}")
                    return False

                success_count = 0
                source_parser_counter: Counter[str] = Counter()
                org_tag_value = org_tag or file_record.org_tag or "DEFAULT"
                is_public_value = (
                    is_public
                    if is_public is not None
                    else (file_record.is_public if file_record.is_public is not None else False)
                )
                await db.execute(ChunkSource.__table__.delete().where(ChunkSource.file_md5 == file_md5))
                await db.execute(DocumentVector.__table__.delete().where(DocumentVector.file_md5 == file_md5))
                es_cleanup_ok = await es_client.delete_by_query(
                    index=search_service.INDEX_NAME,
                    query={"term": {"file_md5": file_md5}},
                )
                if not es_cleanup_ok:
                    raise RuntimeError(f"ES旧文档清理失败: file_md5={file_md5}")

                for chunk, vector in zip(chunks, vectors):
                    if vector is None:
                        continue

                    chunk_meta = chunk.get("metadata", {})
                    chunk_id = chunk["chunk_id"]
                    chunk_text = chunk["text"]
                    chunk_sources = chunk.get("sources") or []

                    db.add(
                        DocumentVector(
                            file_md5=file_md5,
                            chunk_id=chunk_id,
                            text_content=chunk_text,
                            model_version=settings.OPENAI_EMBEDDING_MODEL
                        )
                    )
                    if str(chunk_meta.get("source_parser") or "") in {"vlm_sheet_snapshot", "vlm_diagram"}:
                        logger.info(
                            "VLM_DB_WRITE: table=document_vectors, file_md5=%s, chunk_id=%s, source_parser=%s, source_type=%s, sheet=%s, page=%s, text_content=%s",
                            file_md5,
                            chunk_id,
                            chunk_meta.get("source_parser"),
                            chunk_meta.get("chunk_type"),
                            chunk_meta.get("sheet"),
                            chunk_meta.get("page"),
                            self._clean_text(chunk_text),
                        )

                    es_doc = {
                        "file_md5": file_md5,
                        "chunk_id": chunk_id,
                        "text_content": chunk_text,
                        "vector": vector,
                        "user_id": user_id,
                        "org_tag": org_tag_value,
                        "is_public": is_public_value,
                        "kb_profile": selected_profile,
                        "file_name": file_name,
                        "model_version": settings.OPENAI_EMBEDDING_MODEL,
                        "chunk_type": chunk_meta.get("chunk_type"),
                        "page": chunk_meta.get("page"),
                        "section": chunk_meta.get("section"),
                        "sheet": chunk_meta.get("sheet"),
                        "image_path": chunk_meta.get("image_path"),
                    }
                    doc_id = f"{file_md5}_{chunk_id}"
                    result = await es_client.index_document(
                        index=search_service.INDEX_NAME,
                        document=es_doc,
                        doc_id=doc_id
                    )
                    if not result:
                        raise RuntimeError(f"ES写入失败: file_md5={file_md5}, chunk_id={chunk_id}")
                    success_count += 1

                    if chunk_sources:
                        for idx, src in enumerate(chunk_sources):
                            parser_name = str(src.get("source_parser") or "unknown")
                            if self._is_vlm_parser(parser_name):
                                resolved_image_path = self._resolve_vlm_image_path(src, chunk_meta)
                                if not resolved_image_path:
                                    logger.warning(
                                        "VLM_SOURCE_DROP_NO_IMAGE: file=%s, file_md5=%s, chunk_id=%s, source_order=%s, source_parser=%s, source_type=%s",
                                        file_name,
                                        file_md5,
                                        chunk_id,
                                        idx,
                                        parser_name,
                                        src.get("source_type"),
                                    )
                                    continue
                                src = dict(src)
                                src["image_path"] = resolved_image_path
                            source_parser_counter[parser_name] += 1
                            db.add(
                                ChunkSource(
                                    file_md5=file_md5,
                                    chunk_id=chunk_id,
                                    source_order=idx,
                                    block_index=src.get("block_index"),
                                    source_type=src.get("source_type"),
                                    source_parser=src.get("source_parser"),
                                    sheet=src.get("sheet"),
                                    page=src.get("page"),
                                    section=src.get("section"),
                                    image_path=src.get("image_path"),
                                    text_preview=self._clean_preview_text(
                                        src.get("text_preview"),
                                        src.get("source_type"),
                                    )[:500],
                                )
                            )
                            if parser_name in {"vlm_sheet_snapshot", "vlm_diagram"}:
                                logger.info(
                                    "VLM_DB_WRITE: table=chunk_sources, file_md5=%s, chunk_id=%s, source_order=%s, source_type=%s, source_parser=%s, sheet=%s, page=%s, image_path=%s, text_preview=%s",
                                    file_md5,
                                    chunk_id,
                                    idx,
                                    src.get("source_type"),
                                    src.get("source_parser"),
                                    src.get("sheet"),
                                    src.get("page"),
                                    src.get("image_path"),
                                    self._clean_text(src.get("text_preview")),
                                )
                    else:
                        parser_name = str(chunk_meta.get("source_parser") or "unknown")
                        if self._is_vlm_parser(parser_name):
                            resolved_meta_image = self._clean_text(chunk_meta.get("image_path"))
                            if not resolved_meta_image:
                                logger.warning(
                                    "VLM_FALLBACK_SOURCE_DROP_NO_IMAGE: file=%s, file_md5=%s, chunk_id=%s, source_parser=%s, chunk_type=%s",
                                    file_name,
                                    file_md5,
                                    chunk_id,
                                    parser_name,
                                    chunk_meta.get("chunk_type"),
                                )
                                continue
                            chunk_meta = dict(chunk_meta)
                            chunk_meta["image_path"] = resolved_meta_image
                        source_parser_counter[parser_name] += 1
                        db.add(
                            ChunkSource(
                                file_md5=file_md5,
                                chunk_id=chunk_id,
                                source_order=0,
                                block_index=chunk_meta.get("block_index"),
                                source_type=chunk_meta.get("chunk_type"),
                                source_parser=chunk_meta.get("source_parser"),
                                sheet=chunk_meta.get("sheet"),
                                page=chunk_meta.get("page"),
                                section=chunk_meta.get("section"),
                                image_path=chunk_meta.get("image_path"),
                                text_preview=self._clean_preview_text(
                                    chunk_text,
                                    chunk_meta.get("chunk_type"),
                                )[:500],
                            )
                        )
                        if parser_name in {"vlm_sheet_snapshot", "vlm_diagram"}:
                            logger.info(
                                "VLM_DB_WRITE: table=chunk_sources, file_md5=%s, chunk_id=%s, source_order=%s, source_type=%s, source_parser=%s, sheet=%s, page=%s, image_path=%s, text_preview=%s",
                                file_md5,
                                chunk_id,
                                0,
                                chunk_meta.get("chunk_type"),
                                chunk_meta.get("source_parser"),
                                chunk_meta.get("sheet"),
                                chunk_meta.get("page"),
                                chunk_meta.get("image_path"),
                                self._clean_text(chunk_text),
                            )

                await db.commit()
                await es_client.refresh_index(search_service.INDEX_NAME)
                logger.info(
                    "CHUNK_SOURCE_PARSER_COUNTS: file=%s, file_md5=%s, counts=%s",
                    file_name,
                    file_md5,
                    dict(source_parser_counter),
                )
                logger.info(
                    "VLM_PERSIST_SUMMARY: file=%s, file_md5=%s, vlm_sheet_snapshot=%s, vlm_diagram=%s",
                    file_name,
                    file_md5,
                    source_parser_counter.get("vlm_sheet_snapshot", 0),
                    source_parser_counter.get("vlm_diagram", 0),
                )

                logger.info(
                    f"文档处理完成: file_md5={file_md5}, 成功索引 {success_count}/{len(chunks)} 个文本块"
                )
                return success_count > 0

        except Exception as e:
            try:
                await es_client.delete_by_query(
                    index=search_service.INDEX_NAME,
                    query={"term": {"file_md5": file_md5}},
                )
            except Exception:
                pass
            logger.error(f"处理文档失败: file_md5={file_md5}, 错误: {e}", exc_info=True)
            return False

    async def handle_kafka_message(self, message) -> bool:
        """
        处理Kafka消息
        """
        try:
            message_data = message.value
            if not isinstance(message_data, dict):
                logger.error(f"无效的消息格式: {type(message_data)}, 消息内容: {message_data}")
                return False

            file_md5 = message_data.get("file_md5")
            file_name = message_data.get("file_name")
            storage_path = message_data.get("storage_path")
            user_id = message_data.get("user_id")
            org_tag = message_data.get("org_tag")
            is_public = message_data.get("is_public", False)
            kb_profile = message_data.get("kb_profile")

            missing_fields = []
            if not file_md5:
                missing_fields.append("file_md5")
            if not file_name:
                missing_fields.append("file_name")
            if not storage_path:
                missing_fields.append("storage_path")
            if not user_id:
                missing_fields.append("user_id")
            if missing_fields:
                logger.error(f"消息缺少必要字段: {missing_fields}, 消息内容: {message_data}")
                return False

            logger.info(
                f"收到文档处理消息: file_md5={file_md5}, file_name={file_name}, "
                f"user_id={user_id}, topic={message.topic}, partition={message.partition}, offset={message.offset}"
            )

            lock_key = self._processing_lock_key(file_md5=file_md5, user_id=user_id)
            lock_token = f"{message.topic}:{message.partition}:{message.offset}"

            if await self._has_done_marker(file_md5=file_md5, user_id=user_id):
                logger.info("跳过重复消息（done_marker 命中）: file_md5=%s, user_id=%s", file_md5, user_id)
                return True

            if await self._is_file_done(file_md5=file_md5, user_id=user_id):
                await self._set_done_marker(file_md5=file_md5, user_id=user_id)
                logger.info("跳过重复消息（DB 状态已完成）: file_md5=%s, user_id=%s", file_md5, user_id)
                return True

            acquired = await self._acquire_processing_lock(lock_key=lock_key, token=lock_token)
            if not acquired:
                logger.info("跳过重复消息（文件正在处理中）: file_md5=%s, user_id=%s", file_md5, user_id)
                return True

            await self._mark_file_status(file_md5=file_md5, user_id=user_id, status=FILE_STATUS_PROCESSING)

            max_retries = 3
            retry_count = 0
            success = False

            try:
                while retry_count < max_retries and not success:
                    try:
                        success = await self.process_document(
                            file_md5=file_md5,
                            file_name=file_name,
                            storage_path=storage_path,
                            user_id=user_id,
                            org_tag=org_tag,
                            is_public=is_public,
                            kb_profile=kb_profile,
                        )
                        if success:
                            await self._mark_file_status(
                                file_md5=file_md5, user_id=user_id, status=FILE_STATUS_DONE
                            )
                            await self._set_done_marker(file_md5=file_md5, user_id=user_id)
                            logger.info(f"文档处理成功: file_md5={file_md5}, 重试次数: {retry_count}")
                            break

                        retry_count += 1
                        if retry_count < max_retries:
                            logger.warning(f"文档处理失败，准备重试 ({retry_count}/{max_retries}): file_md5={file_md5}")
                            await asyncio.sleep(2 ** retry_count)
                        else:
                            logger.error(f"文档处理失败，已达到最大重试次数: file_md5={file_md5}")

                    except Exception as e:
                        retry_count += 1
                        logger.error(
                            f"文档处理异常 (重试 {retry_count}/{max_retries}): file_md5={file_md5}, 错误: {e}",
                            exc_info=True,
                        )
                        if retry_count < max_retries:
                            await asyncio.sleep(2 ** retry_count)
                        else:
                            logger.error(f"文档处理异常，已达到最大重试次数: file_md5={file_md5}")

                if not success:
                    await self._mark_file_status(file_md5=file_md5, user_id=user_id, status=FILE_STATUS_FAILED)
                    await self._publish_to_dlq(
                        message_data=message_data,
                        message=message,
                        fail_reason=f"process_document failed after retries: file_md5={file_md5}",
                    )
                    return True
                return True
            finally:
                await self._release_processing_lock(lock_key=lock_key, token=lock_token)

        except Exception as e:
            logger.error(f"处理Kafka消息失败: {e}", exc_info=True)
            return False


document_processor_service = DocumentProcessorService()
